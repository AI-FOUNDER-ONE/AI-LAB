import docx
import zipfile
import io
from PIL import Image, ImageDraw
import os

def create_test_docx(filename):
    doc = docx.Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph with some text.')
    
    # Create dummy image
    img = Image.new('RGB', (60, 30), color = 'red')
    img_path = "temp_img.png"
    img.save(img_path)
    
    doc.add_picture(img_path, width=docx.shared.Inches(1.0))
    doc.save(filename)
    
    # Clean up image
    os.remove(img_path)
    print(f"Created {filename}")

def test_parse_docx(path):
    print(f"Parsing {path}...")
    text = ""
    images = []
    try:
        doc = docx.Document(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        with zipfile.ZipFile(path) as z:
            for name in z.namelist():
                if name.startswith("word/media/") and name.lower().endswith(('.png', '.jpg', '.jpeg')):
                    print(f"Found image: {name}")
                    img_data = z.read(name)
                    try:
                        img = Image.open(io.BytesIO(img_data))
                        images.append(img)
                    except Exception as e:
                        print(f"Error opening image: {e}")
    except Exception as e:
        print(f"Error parsing: {e}")
        
    print(f"Extracted Text Length: {len(text)}")
    print(f"Extracted Text: {text[:50]}...")
    print(f"Extracted Images Count: {len(images)}")
    
    if len(text) > 0 and len(images) > 0:
        print("PASS: Text and Image extracted.")
    else:
        print("FAIL: Missing text or image.")

if __name__ == "__main__":
    docx_name = "test_doc_vision.docx"
    create_test_docx(docx_name)
    test_parse_docx(docx_name)
    # Cleanup docx
    if os.path.exists(docx_name):
        os.remove(docx_name)
