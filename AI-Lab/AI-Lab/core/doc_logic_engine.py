"""
core/doc_logic_engine.py - 文档逻辑架构引擎
=========================================

该模块实现了 AI 自动化文档克隆与内容重组的核心逻辑。
主要功能：
1. 原型文档结构分析 (Prototype Analysis)
2. 新内容原子化提取 (Atom Extraction)
3. 逻辑映射与重组 (Mapping & Restructuring)

该模块设计为与 LLM 提供商无关，接受标准的 OpenAI 兼容客户端。
"""

import json
import logging
from typing import Dict, List, Any, Optional

# 配置日志
logger = logging.getLogger(__name__)

class DocumentLogicEngine:
    """文档逻辑架构引擎"""

    def __init__(self, client: Any, model: str):
        """
        初始化引擎。

        Args:
            client: 兼容 OpenAI 接口的客户端对象 (需有 chat.completions.create 方法)
            model: 使用的模型名称 (e.g., "gpt-4", "qwen-max")
        """
        self.client = client
        self.model = model

    def analyze_prototype(self, prototype_text: str) -> Dict[str, Any]:
        """
        [Step 1] 分析原型文档，提取逻辑块结构。
        
        Args:
             prototype_text: 原型文档的纯文本内容
             
        Returns:
            JSON 结构描述文档的逻辑骨架
        """
        system_prompt = """你是一个文档结构分析专家。
请分析用户提供的文档内容，提取其核心逻辑结构（骨架）。
输出必须是 JSON 格式，包含一个有序的 "sections" 列表，每个 section 包含 "title" 和 "description"（该块通常包含什么内容）。
例如：
{
  "sections": [
    {"title": "背景介绍", "description": "项目发起的缘由和背景"},
    {"title": "现状痛点", "description": "当前存在的问题和挑战"}
  ]
}
"""
        return self._call_llm(system_prompt, prototype_text, json_mode=True)

    def extract_atoms(self, new_content: str) -> List[Dict[str, Any]]:
        """
        [Step 2] 对新内容进行“原子化提取”。
        
        Args:
            new_content: 用户提供的杂乱新内容
            
        Returns:
            原子化信息列表
        """
        system_prompt = """你是一个信息抽取专家。
请将输入的内容拆解为“原子化”的信息点。识别哪些是核心目标、技术指标、商务条款等。
输出必须是 JSON 格式，包含 "atoms" 列表。
"""
        return self._call_llm(system_prompt, new_content, json_mode=True)

    def process_content_to_template(self, raw_new_content: str, prototype_structure: Dict[str, Any]) -> Dict[str, Any]:
        """
        [Step 3 (Core)] 将新内容映射到旧结构，并进行风格化重写。
        
        Args:
            raw_new_content: 原始新内容字符串
            prototype_structure: analyze_prototype 返回的结构字典
            
        Returns:
            生成的最终文档结构 JSON
        """
        system_prompt = """
你是一个文档重构专家。你的任务是：
1. 分析原型文档的 JSON 结构模板。
2. 深度解析用户提供的新内容，将其分类、拆解。
3. 将拆解后的内容映射到模板的对应 Key 中。
4. 润色：使用原型文档的语言风格（正式度、行业术语、句式习惯）重写新内容。
5. 自动补全：如果新内容缺失原型中的某个必要模块，请基于上下文自动补全该模块。
6. 输出：仅输出合法的 JSON 格式数据。结构必须严格遵循原型结构。
"""
        
        user_prompt = f"""
原型结构模板：
{json.dumps(prototype_structure, ensure_ascii=False, indent=2)}

新内容原始信息：
{raw_new_content}

请生成符合上述结构的 JSON 内容。
"""
        
        # 这里使用较大的 max_tokens 以容纳生成的文档
        return self._call_llm(system_prompt, user_prompt, json_mode=True, max_tokens=4000)

    def _call_llm(self, system_prompt: str, user_prompt: str, json_mode: bool = False, max_tokens: int = 2000) -> Any:
        """内部工具函数：调用 LLM"""
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=0.3, # 保持结构化输出的稳定性
                max_tokens=max_tokens,
                response_format={"type": "json_object"} if json_mode else None
            )
            content = response.choices[0].message.content
            if json_mode:
                # 增强：尝试清理 Markdown 代码块
                cleaned_content = content.strip()
                if cleaned_content.startswith("```"):
                    # 去掉第一行 (```json) 和最后一行 (```)
                    lines = cleaned_content.splitlines()
                    if len(lines) >= 2:
                        # 查找第一个和最后一个 ``` 的位置
                        import re
                        match = re.search(r"```(json)?\s*([\s\S]*?)\s*```", cleaned_content)
                        if match:
                            cleaned_content = match.group(2)
                        else:
                            # 简单的去头去尾 fallback
                            cleaned_content = "\n".join(lines[1:-1])
                
                # 尝试解析
                try:
                    return json.loads(cleaned_content)
                except json.JSONDecodeError:
                    # 如果清理后仍失败，尝试在整个内容中寻找 JSON 对象
                    import re
                    json_match = re.search(r"\{[\s\S]*\}", content)
                    if json_match:
                        return json.loads(json_match.group(0))
                    raise # 重新抛出异常
                        
            return content
        except Exception as e:
            logger.error(f"LLM Call Failed: {e}")
            raise RuntimeError(f"LLM 调用失败: {str(e)}")

    def read_docx(self, file_path: str) -> str:
        """读取 Word 文档 (.docx) 文本内容"""
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
        except Exception as e:
            raise RuntimeError(f"读取 Word 文档失败: {str(e)}")

    def save_to_docx(self, json_data: Dict[str, Any], output_path: str):
        """将生成的 JSON 数据保存为 Word 文档
        
        预期 JSON 结构:
        {
            "sections": [
                {"title": "...", "content": "..."}
            ]
        }
        """
        try:
            import docx
            doc = docx.Document()
            
            # 假设最外层结构
            sections = json_data.get("sections", [])
            
            # 如果不是标准结构，尝试智能解析
            if not sections and isinstance(json_data, dict):
                # 可能是扁平字典，Key作为标题
                for key, value in json_data.items():
                    doc.add_heading(key, level=1)
                    if isinstance(value, str):
                        doc.add_paragraph(value)
                    elif isinstance(value, list) or isinstance(value, dict):
                        doc.add_paragraph(json.dumps(value, ensure_ascii=False, indent=2))
            else:
                for sec in sections:
                    title = sec.get("title", "未命名章节")
                    content = sec.get("content", "")
                    doc.add_heading(title, level=1)
                    doc.add_paragraph(str(content)) # 确保转为字符串
            
            doc.save(output_path)
            logger.info(f"文档已保存至: {output_path}")
            
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
        except Exception as e:
            logger.error(f"保存 Word 文档失败: {str(e)}")
            raise RuntimeError(f"保存 Word 文档失败: {str(e)}")

