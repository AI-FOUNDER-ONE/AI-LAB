"""
commander_crew_v2.py - CommanderCrew 增强型多 Agent 辩论引擎 V2
================================================================

核心升级:
  1. PersonalityEngine:  Agent 个性化风格引擎 (随机化发言风格)
  2. WarRoomContext:      增强消息总线 (相关历史检索、结构化投票)
  3. DebateRouter:        智能路由器 (语义相关性、@提及优先、冲突触发)
  4. DebateWorker:        多轮辩论引擎 (结构化投票、子辩论、历史引用)
  5. CommanderCrew:       主编排器 (用户干预广播、PM 即时响应)

工作流:
  Phase 1: CKO Grounding (需求访谈)
  Phase 2: Debate Loop (PM主持 → 多轮博弈 → 结构化投票 → 共识/子辩论)
  Phase 3: Production (Coder 编码 + Tester 反馈循环)
  Phase 4: PM Final Review (最终审批)
"""

import os
import re
import sys
import random
import threading
import time
from typing import List, Dict, Optional, Any, Tuple
from dataclasses import dataclass, field
from collections import deque

from PyQt6.QtCore import QObject, QThread, pyqtSignal

from crewai import Crew, Process, Task, Agent
from config import API_KEYS, AppState

# ============================================================
#  环境设置
# ============================================================
os.environ["OPENAI_API_KEY"] = API_KEYS.get("hiapi", "")
os.environ["OPENAI_API_BASE"] = "https://hiapi.online/v1"

from agents.crew_agents import create_agents
from core.crew_tasks import create_grounding_task

# Windows UTF-8 编码修复
if sys.platform == "win32":
    try:
        import codecs
        if sys.stdout.encoding != 'utf-8':
            sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())
            sys.stderr = codecs.getwriter("utf-8")(sys.stderr.detach())
    except Exception:
        pass

# ============================================================
#  配置常量
# ============================================================
DEBATE_MAX_ROUNDS = 5
DEBATE_MIN_ROUNDS = 2
CONSENSUS_THRESHOLD = 0.7
PRODUCTION_MAX_LOOPS = 3
THINKING_DELAY_MIN_MS = 200   # 最小思考延迟
THINKING_DELAY_MAX_MS = 800   # 最大思考延迟

# 冲突 / 同意检测关键词
CONFLICT_KEYWORDS_CN = [
    "反对", "不同意", "不行", "有问题", "质疑", "担心",
    "但是", "然而", "存在风险", "不可行", "重新考虑",
    "不够", "缺陷", "漏洞", "我认为不",
]
CONFLICT_KEYWORDS_EN = [
    "disagree", "however", "but", "issue", "problem",
    "concern", "risk", "alternative", "instead",
]
AGREEMENT_KEYWORDS = [
    "同意", "赞成", "认可", "支持", "没问题", "可以",
    "LGTM", "agree", "approved", "好的方案", "我同意",
    "通过", "确认", "接受",
]

# Agent 专长关键词 — 用于语义相关性路由
AGENT_EXPERTISE = {
    "CKO": ["需求", "目标", "用户", "场景", "协议", "任务书", "需求分析"],
    "PM": ["进度", "优先级", "风险", "里程碑", "计划", "资源", "管理"],
    "Arch": ["架构", "技术栈", "模块", "API", "数据库", "系统设计", "组件",
             "性能", "扩展性", "微服务", "scalability"],
    "Designer": ["UI", "UX", "界面", "交互", "原型", "设计", "用户体验",
                 "布局", "色彩", "实现方案", "详细设计", "美观", "好看", "风格"],
    "Coder": ["代码", "实现", "函数", "类", "算法", "Python", "编码",
              "bug", "重构", "测试用例"],
    "Tester": ["测试", "验证", "质量", "缺陷", "回归", "覆盖率",
               "验收", "边界条件", "异常处理", "BUG", "报错", "失败"],
}


# ==============================================================================
#  1. PersonalityEngine — Agent 个性化风格引擎
# ==============================================================================

class PersonalityEngine:
    """为每个 Agent 生成随机化的发言风格指令。

    每轮辩论随机选取一种风格，注入到 Prompt 中，
    使 Agent 回复呈现不同的语气和详细程度。
    """

    # 每个角色的风格池
    STYLE_POOL = {
        "CKO": [
            "用学术导师口吻，温和但严谨，多用反问引导思考",
            "用简洁的列表形式总结要点，不超过 5 条",
            "用类比方式解释复杂概念，引用行业案例",
        ],
        "PM": [
            "用指挥官口吻，果断决策，明确下达行动项",
            "用协作者口吻，先肯定各方贡献再指出改进方向",
            "用商业分析师口吻，从 ROI 和风险角度评估方案",
        ],
        "Arch": [
            "用技术极客风格，深入底层原理，引用技术规范",
            "用实用主义风格，聚焦可落地性和工程成本",
            "用系统思维风格，从全局视角分析模块间耦合关系",
        ],
        "Designer": [
            "用用户体验专家风格，始终从终端用户视角出发",
            "用工程化设计师风格，注重方案的可实现性和技术约束",
            "用创新型设计师风格，提出大胆的创新方案并分析可行性",
        ],
        "Coder": [
            "用资深工程师风格，注重代码质量和最佳实践",
            "用全栈开发者风格，同时考虑前后端和部署流程",
            "用效率至上风格，追求最简方案和快速迭代",
        ],
        "Tester": [
            "用严格质检官风格，对每个细节都不放过",
            "用风险分析师风格，聚焦高优先级的潜在失败点",
            "用建设性评审员风格，发现问题的同时给出修复建议",
        ],
    }

    @classmethod
    def get_style_directive(cls, role: str) -> str:
        """随机选取一种发言风格指令。"""
        pool = cls.STYLE_POOL.get(role, ["保持专业和简洁"])
        return random.choice(pool)

    @staticmethod
    def thinking_delay():
        """随机思考延迟 (200-800ms)，模拟真实讨论节奏。"""
        delay = random.randint(THINKING_DELAY_MIN_MS, THINKING_DELAY_MAX_MS)
        time.sleep(delay / 1000.0)


# ==============================================================================
#  2. WarRoomMessage — 增强消息数据类
# ==============================================================================

@dataclass
class WarRoomMessage:
    """War Room 中的一条消息，包含元数据用于辩论分析。"""
    role: str
    content: str
    intent: str = "statement"
    mentions: List[str] = field(default_factory=list)
    sentiment: str = "neutral"
    round_num: int = 0
    timestamp: int = 0
    priority: int = 0  # 0=普通, 10=用户干预(最高)

    def __post_init__(self):
        """自动检测意图、情感和 @提及。"""
        self.mentions = re.findall(r"@(\w+)", self.content)
        content_lower = self.content.lower()

        if self.role == "Commander":
            self.intent = "command"
            self.priority = 10
        elif any(kw in self.content for kw in CONFLICT_KEYWORDS_CN):
            self.intent = "critique"
            self.sentiment = "negative"
        elif any(kw in content_lower for kw in CONFLICT_KEYWORDS_EN):
            self.intent = "critique"
            self.sentiment = "negative"
        elif any(kw in self.content for kw in AGREEMENT_KEYWORDS):
            self.intent = "agreement"
            self.sentiment = "positive"
        elif "?" in self.content or "？" in self.content:
            self.intent = "question"
        else:
            self.intent = "statement"
            self.sentiment = "neutral"


# ==============================================================================
#  3. WarRoomContext — 增强消息总线 & 共识追踪
# ==============================================================================

class WarRoomContext:
    """War Room 会话上下文管理器。

    管理完整的辩论历史、Agent 立场追踪、结构化投票、相关历史检索。
    """

    def __init__(self):
        self.history: List[WarRoomMessage] = []
        self.artifacts: Dict[str, Any] = {}
        self.mission_protocol: Dict[str, Any] = {}
        self.task_type: str = "SOFTWARE"

        # 共识追踪
        self.consensus_tracker: Dict[str, str] = {
            "PM": "neutral", "Arch": "neutral", "Designer": "neutral",
            "Coder": "neutral", "Tester": "neutral",
        }
        # 结构化投票记录: [{role, stance, reason, round}]
        self.vote_records: List[Dict[str, Any]] = []
        # 冲突记录
        self.active_conflicts: List[Dict[str, Any]] = []
        # [NEW] 结构化约束（从 CKO 输出中提取的死命令）
        self.constraints: List[str] = []

    def add_message(self, role: str, content: str, round_num: int = 0,
                    priority: int = 0) -> WarRoomMessage:
        """添加消息并自动分析意图和情感。"""
        msg = WarRoomMessage(
            role=role, content=content, round_num=round_num,
            timestamp=len(self.history) + 1, priority=priority
        )
        self.history.append(msg)

        # 自动更新共识追踪
        if role in self.consensus_tracker:
            if msg.intent == "agreement":
                self.consensus_tracker[role] = "agree"
            elif msg.intent == "critique":
                self.consensus_tracker[role] = "disagree"
        return msg

    def get_recent_history(self, limit: int = 8) -> str:
        """格式化最近消息供 LLM 上下文注入。"""
        # 兼容旧代码，内部调用 get_safe_history
        return self.get_safe_history(limit=limit)

    def get_safe_history(self, limit: int = 15, max_chars: int = 12000) -> str:
        """智能获取历史记录，防止 Token 溢出。

        策略:
        1. 总是保留 CKO 的锚点 (如果存在)。
        2. 保留最近的 N 条消息。
        3. 如果总字数超过 max_chars，则仅保留最近的 3 条完整消息，中间部分做摘要或截断。
        """
        if not self.history:
            return "（暂无讨论历史）"

        # 1. 获取目标消息片段
        target_msgs = self.history[-limit:]
        
        # 2. 估算字符数
        total_chars = sum(len(m.content) for m in target_msgs)
        
        lines = []
        # 如果未超标，正常返回
        if total_chars <= max_chars:
            for msg in target_msgs:
                intent_tag = f"[{msg.intent}]" if msg.intent != "statement" else ""
                prio_tag = "🔴 " if msg.priority >= 10 else ""
                lines.append(f"{prio_tag}**{msg.role}** {intent_tag}: {msg.content}")
        else:
            # 3. 超标处理: 保留锚点 + 最近 3 条
            lines.append(f"【系统提示】由于历史记录过长（{total_chars} 字符），已自动隐藏中间部分...\n")
            
            # 确保至少显示最近 3 条
            safe_recent = target_msgs[-3:]
            for msg in safe_recent:
               intent_tag = f"[{msg.intent}]" if msg.intent != "statement" else ""
               prio_tag = "🔴 " if msg.priority >= 10 else ""
               lines.append(f"{prio_tag}**{msg.role}** {intent_tag}: {msg.content}")

        return "\n\n".join(lines)

    def get_relevant_history(self, role: str, limit: int = 3) -> str:
        """根据角色专长检索最相关的历史消息（非自己的发言）。

        利用关键词交集计算相关性，返回最相关的 N 条消息的引文。
        """
        expertise = AGENT_EXPERTISE.get(role, [])
        if not expertise or not self.history:
            return ""

        scored: List[Tuple[float, WarRoomMessage]] = []
        for msg in self.history:
            if msg.role == role:  # 跳过自己的消息
                continue
            # 关键词匹配相关性
            score = sum(1 for kw in expertise if kw in msg.content)
            if score > 0:
                scored.append((score, msg))

        # 按相关性排序，取 top N
        scored.sort(key=lambda x: x[0], reverse=True)
        top = scored[:limit]

        if not top:
            return ""

        lines = ["[相关历史引文]"]
        for _, msg in top:
            preview = msg.content[:120] + ("..." if len(msg.content) > 120 else "")
            lines.append(f"- {msg.role}(第{msg.round_num}轮): \"{preview}\"")
        return "\n".join(lines)

    def get_stance_summary(self) -> str:
        """生成各 Agent 的立场态势图。"""
        agrees = [r for r, s in self.consensus_tracker.items() if s == "agree"]
        disagrees = [r for r, s in self.consensus_tracker.items() if s == "disagree"]
        neutrals = [r for r, s in self.consensus_tracker.items() if s == "neutral"]

        lines = ["## 当前立场态势"]
        if agrees:
            lines.append(f"- ✅ 同意方: {', '.join(agrees)}")
        if disagrees:
            lines.append(f"- ❌ 反对方: {', '.join(disagrees)}")
        if neutrals:
            lines.append(f"- ⏳ 中立方: {', '.join(neutrals)}")
        if self.active_conflicts:
            lines.append("\n### 活跃分歧:")
            for c in self.active_conflicts[-3:]:
                lines.append(f"- {c.get('description', '未指定')}")
        return "\n".join(lines)

    def register_vote(self, role: str, stance: str, reason: str,
                      round_num: int = 0):
        """注册结构化投票。"""
        self.vote_records.append({
            "role": role, "stance": stance,
            "reason": reason, "round": round_num
        })
        if stance in ("agree", "conditional"):
            self.consensus_tracker[role] = "agree"
        elif stance == "disagree":
            self.consensus_tracker[role] = "disagree"

    def check_consensus(self) -> bool:
        """检查是否达成共识。"""
        total = len(self.consensus_tracker)
        agrees = sum(1 for s in self.consensus_tracker.values() if s == "agree")
        return (agrees / total if total > 0 else 0) >= CONSENSUS_THRESHOLD

    def get_last_message(self) -> Optional[WarRoomMessage]:
        return self.history[-1] if self.history else None

    def get_mission_anchor(self) -> str:
        """获取原始需求锚点（CKO 的第一次发言或用户输入）。

        [V2] 取消硬编码截断，改为智能压缩：
        - 上限 2000 字符
        - 超长时保留 头500 + 尾1000（尾部含约束/负面清单），中间省略
        """
        # 优先找 CKO 的第一条完整分析
        for msg in self.history:
            if msg.role == "CKO":
                return self._smart_compress_anchor(msg.content, max_len=2000)

        # 如果没找到 CKO (e.g. bypass mode), 返回最早的历史消息
        if self.history:
            return self._smart_compress_anchor(self.history[0].content, max_len=1500)

        return "【🛡️ 原始需求锚点】(暂无)"

    def _smart_compress_anchor(self, content: str, max_len: int = 2000) -> str:
        """智能压缩锚点内容。

        策略：
        - 内容 <= max_len：完整保留
        - 内容 > max_len：保留头 500 字 + 尾 1000 字（尾部包含约束条件和负面清单），
          中间用 [...已省略 N 字...] 替代
        """
        if len(content) <= max_len:
            return f"【🛡️ 原始需求锚点】\n{content}"

        # 头部 500 字 + 尾部 1000 字（尾部含约束/负面清单，权重更高）
        head_len = 500
        tail_len = 1000
        omitted = len(content) - head_len - tail_len
        compressed = (
            f"{content[:head_len]}\n"
            f"\n[...已省略 {omitted} 字...]\n\n"
            f"{content[-tail_len:]}"
        )
        return f"【🛡️ 原始需求锚点】\n{compressed}"

    def extract_constraints(self, cko_output: str) -> List[str]:
        """从 CKO 输出中提取结构化约束。

        匹配模式：
        - 标签格式：【约束】、【技术栈】、【禁止】、【平台】、【约束清单】
        - 列表格式：- 技术栈: xxx  /  - 禁止: xxx
        提取后存储到 self.constraints
        """
        constraints = []

        # 模式 1：匹配 【约束清单】 区块（CKO 增强输出）
        constraint_block = re.search(
            r"【约束清单】[\s\S]*?(?=\n【|$)", cko_output
        )
        if constraint_block:
            block_text = constraint_block.group(0)
            # 提取区块中每一行非空内容
            for line in block_text.split("\n"):
                line = line.strip()
                if line and line != "【约束清单】":
                    # 去除列表符号
                    line = re.sub(r"^[\-\*\d+\.\s]+", "", line).strip()
                    if line:
                        constraints.append(line)

        # 模式 2：匹配独立标签行 【约束】xxx、【技术栈】xxx、【禁止】xxx、【平台】xxx
        tag_patterns = re.findall(
            r"【(约束|技术栈|禁止|平台|限制|必须|不得)】\s*[:：]?\s*(.+)",
            cko_output
        )
        for tag, value in tag_patterns:
            constraint_text = f"{tag}: {value.strip()}"
            if constraint_text not in constraints:
                constraints.append(constraint_text)

        # 模式 3：匹配常见列表格式 (e.g. - 技术栈: Python + FastAPI)
        list_patterns = re.findall(
            r"[-\*]\s*(技术栈|平台|禁止|约束|限制|必须使用|不得使用)\s*[:：]\s*(.+)",
            cko_output
        )
        for tag, value in list_patterns:
            constraint_text = f"{tag}: {value.strip()}"
            if constraint_text not in constraints:
                constraints.append(constraint_text)

        # 存储到实例
        self.constraints = constraints
        if constraints:
            print(f"[CONSTRAINTS] ✅ 提取到 {len(constraints)} 条结构化约束:")
            for i, c in enumerate(constraints, 1):
                print(f"  {i}. {c}")
        else:
            print("[CONSTRAINTS] ⚠️ 未提取到结构化约束")

        return constraints

    def reset(self):
        """重置上下文。"""
        self.history.clear()
        self.artifacts.clear()
        self.mission_protocol.clear()
        self.task_type = "SOFTWARE"
        self.vote_records.clear()
        for role in self.consensus_tracker:
            self.consensus_tracker[role] = "neutral"
        self.active_conflicts.clear()
        self.constraints.clear()  # [NEW] 清空约束


# ==============================================================================
#  4. DebateRouter — 增强智能路由器
# ==============================================================================

class DebateRouter:
    """辩论路由器 — 决定谁是下一个发言人，基于语义相关性检测。

    路由优先级:
      1. 用户干预 — Commander 消息后 PM 立即响应
      2. @提及 — 被点名的 Agent 最优先
      3. 语义相关性 — 消息内容与 Agent 专长的关键词匹配
      4. 阶段默认路由 — 根据 task_type 决定辩论顺序
    """

    ALL_ROLES = ["CKO", "PM", "Arch", "Designer", "Coder", "Tester"]

    def __init__(self, ctx: WarRoomContext):
        self.ctx = ctx

    def get_debate_order(self, round_num: int) -> List[str]:
        """根据任务类型和轮次返回辩论发言顺序。"""
        task_type = self.ctx.task_type

        if task_type == "RESEARCH":
            base_order = ["CKO", "Arch", "Designer"]
        elif task_type == "DESIGN":
            base_order = ["Designer", "Arch", "CKO"]
        elif task_type == "ENGINEERING":
            base_order = ["Arch", "Designer", "Coder"]
        else:
            base_order = ["Arch", "Designer", "Coder"]

        # 后续轮次：让反对者先说
        if round_num > 2:
            disagrees = [r for r in base_order
                         if self.ctx.consensus_tracker.get(r) == "disagree"]
            others = [r for r in base_order if r not in disagrees]
            base_order = disagrees + others
        return base_order

    def get_production_order(self) -> List[str]:
        return ["Coder", "Tester"]

    def should_agent_respond(self, agent_role: str,
                             latest_msg: WarRoomMessage) -> bool:
        """基于语义相关性判断 Agent 是否应主动发言。

        计算消息内容与 Agent 专长关键词的交集数量。
        """
        if agent_role == latest_msg.role:
            return False  # 不回复自己
        # @mention always responds
        if agent_role.lower() in [m.lower() for m in latest_msg.mentions]:
            return True

        # Keyword matching
        expertise = AGENT_EXPERTISE.get(agent_role, [])
        hits = sum(1 for kw in expertise if kw in latest_msg.content)
        # [OPTIMIZATION] 降低阈值，避免关键角色缺席
        return hits >= 2

    def detect_conflict(self) -> bool:
        """检测最近消息中是否存在冲突。"""
        recent = self.ctx.history[-3:] if len(self.ctx.history) >= 3 else self.ctx.history
        for msg in recent:
            if msg.intent == "critique" and msg.sentiment == "negative":
                conflict = {
                    "instigator": msg.role,
                    "description": f"{msg.role} 对当前方案提出质疑",
                    "content_preview": msg.content[:100],
                }
                if not any(c["instigator"] == msg.role
                           for c in self.ctx.active_conflicts):
                    self.ctx.active_conflicts.append(conflict)
                return True
        return False

    def should_trigger_sub_debate(self) -> bool:
        """是否应触发子辩论（2+ Agent 反对）。"""
        disagrees = sum(1 for s in self.ctx.consensus_tracker.values()
                        if s == "disagree")
        return disagrees >= 2

    def get_response_to_mention(self, mentions: List[str]) -> Optional[str]:
        """解析 @提及，返回应该回应的角色。"""
        for mention in mentions:
            for role in self.ALL_ROLES:
                if role.lower() == mention.lower():
                    return role
        return None





# ==============================================================================
#  6. DebateWorker — 增强型辩论工作线程
# ==============================================================================

class DebateWorker(QThread):
    """在后台线程中运行多轮辩论引擎。

    增强特性:
      - PersonalityEngine 随机化发言风格
      - 结构化投票 (每轮结束后 PM 发起投票)
      - 语义相关性路由 (Agent 主动响应相关话题)
      - 深度历史引用 (Prompt 注入相关历史引文)
      - 子辩论 (严重冲突时触发独立解决)
    """
    agent_responded = pyqtSignal(str, str)    # role, content
    round_completed = pyqtSignal(int)         # round_num
    debate_finished = pyqtSignal(str)         # summary
    production_finished = pyqtSignal(str)     # result
    error_occurred = pyqtSignal(str)          # error
    thought_stream = pyqtSignal(str)          # raw output

    def __init__(self, mode: str, ctx: WarRoomContext, router: DebateRouter,
                 user_input: str = "", parent=None):
        super().__init__(parent)
        self.mode = mode
        self.ctx = ctx
        self.router = router
        self.user_input = user_input
        self._stop_flag = False

    def stop(self):
        self._stop_flag = True
        # [OPTIMIZATION] 立即反馈停止状态
        self.agent_responded.emit("System", "🛑 正在停止中，请等待当前步骤完成...")

    def run(self):
        """线程主逻辑。"""
        # [NEW] 动态创建 Agents，并绑定回调
        self.agents = create_agents(step_callback=self._on_step_callback)
        try:
            if self.mode == "grounding":
                self._run_grounding()
            elif self.mode == "debate":
                self._run_debate()
            elif self.mode == "production":
                self._run_production()
            elif self.mode == "cloning":
                self._run_cloning()
            elif self.mode == "review":
                self._run_review()
        except Exception as e:
            import traceback
            self.error_occurred.emit(
                f"[ERROR] {self.mode}: {e}\n{traceback.format_exc()}")
            
    def _on_step_callback(self, step_output):
        """CrewAI 步骤回调 (Thread-Safe Signal Emit)。"""
        # step_output 是一个 AgentStep 对象，包含 thought, result 等
        if hasattr(step_output, 'thought') and step_output.thought:
            # 清理一下多余的换行
            thought = step_output.thought.strip()
            if thought:
                self.thought_stream.emit(f"{thought}\n")
        
        # 如果有工具调用输出，也可以在这里捕获 (可选)
        # if hasattr(step_output, 'tool_name'): ...

    # ------------------------------------------------------------------
    #  Phase 1: CKO Grounding
    # ------------------------------------------------------------------
    def _run_grounding(self):
        """CKO 需求分析（含自动文件解析）。
        
        增强: 自动检测 [ATTACHMENT: path] 标签，前置解析文档内容
        注入到 CKO Prompt，确保 AI 能读取文件的所有内容。
        """
        user_text = self.user_input
        doc_content = ""
        
        # ---- 自动检测并解析附件 ----
        import re as _re
        attachment_match = _re.search(r"\[ATTACHMENT:\s*(.*?)\]", user_text)
        if attachment_match:
            file_path = attachment_match.group(1).strip()
            # 从消息中移除标签，保留用户的实际文字
            user_text = user_text.replace(attachment_match.group(0), "").strip()
            
            try:
                from tools.document_parser import parse_document
                parsed = parse_document(file_path)
                doc_content = parsed.to_prompt_text()
                
                # [FIX] Safe Print for Windows GBK consoles
                msg = (f"[GROUNDING] ✅ 已自动解析附件: {parsed.filename} "
                       f"({len(doc_content)} 字符, "
                       f"{len(parsed.tables)} 表格, "
                       f"{len(parsed.images_base64)} 图片)")
                try:
                    print(msg)
                except UnicodeEncodeError:
                    # Fallback to ASCII
                    print(msg.replace("✅", "[OK]").replace("❌", "[ERR]").encode('gbk', 'ignore').decode('gbk'))

            except ImportError:
                # 降级: 简单读取
                try:
                    import docx
                    doc = docx.Document(file_path)
                    doc_content = "\n".join(
                        p.text for p in doc.paragraphs if p.text.strip()
                    )
                    print(f"[GROUNDING] [WARN] 使用降级模式解析附件")
                except Exception as e:
                    doc_content = f"[附件解析失败: {e}]"
                    print(f"[GROUNDING] [ERR] 附件解析失败: {e}")
            except Exception as e:
                doc_content = f"[附件解析失败: {e}]"
                print(f"[GROUNDING] [ERR] 附件解析失败: {e}")
        
        # ---- 构建增强 Prompt ----
        attachment_section = ""
        if doc_content:
            attachment_section = (
                f"\n\n{'='*60}\n"
                f"【用户上传的文档内容 — 请仔细阅读并分析】\n"
                f"{'='*60}\n"
                f"{doc_content}\n"
                f"{'='*60}\n"
            )
        
        # [MODIFIED] 使用工厂创建任务 (必须在赋值 description 前创建)
        grounding_task = create_grounding_task(self.agents["CKO"])
        
        grounding_task.description = (
            f"用户输入: {user_text}\n"
            f"{attachment_section}\n"
            f"【任务指令】作为 CKO（首席知识官），请对用户输入进行需求分析：\n"
            f"1. 如果用户上传了文档，请先完整阅读文档内容（包括表格、页眉、批注等），"
            f"然后结合用户的文字说明进行分析\n"
            f"2. 如果用户输入仅为简短问候或模糊描述，你必须：\n"
            f"   - 先礼貌回应\n"
            f"   - 然后提出 3-5 个编号的结构化追问\n"
            f"   - 明确告知用户需要回答这些问题后才能生成任务协议\n"
            f"3. 如果用户已提供具体需求，请直接生成结构化任务协议\n"
            f"4. 输出必须使用简体中文\n"
            f"5. 先写自然语言摘要，再用编号列表提问\n"
            f"6. 请务必在回复末尾用独立一行标识任务类型（三选一）：\n"
            f"   【任务类型】RESEARCH    (调研/文档/策划/文案)\n"
            f"   【任务类型】DESIGN      (设计/绘图/UI/UX)\n"
            f"   【任务类型】ENGINEERING (开发/代码/架构/测试)\n"
            f"   【任务类型】CLONING     (文档克隆/复刻/重组)\n"
            f"7. 【重要】请在回复末尾用【约束清单】标签列出本次任务的所有硬约束，例如：\n"
            f"   【约束清单】\n"
            f"   - 技术栈: xxx\n"
            f"   - 平台: xxx\n"
            f"   - 禁止: xxx\n"
            f"   (如用户未指定硬约束，可填写 '约束清单: 无明确约束')"
        )
        
        crew = Crew(agents=[self.agents["CKO"]], tasks=[grounding_task],
                    process=Process.sequential, verbose=True)
        result = crew.kickoff()
        response = str(result)

        # ---- 解析任务类型 ----
        import re as _re
        type_match = _re.search(r"【任务类型】\s*(RESEARCH|DESIGN|ENGINEERING|CLONING)", response)
        if type_match:
            found_type = type_match.group(1).upper()
            self.ctx.task_type = found_type
            print(f"[GROUNDING] ✅ 识别任务类型: {self.ctx.task_type}")
        else:
            # 默认回退
            self.ctx.task_type = "ENGINEERING"
            print(f"[GROUNDING] ⚠️ 未识别任务类型，默认使用: {self.ctx.task_type}")

        # [NEW] 从 CKO 输出中提取结构化约束（修复辩论跑题核心逻辑）
        self.ctx.extract_constraints(response)

        # 关键: 发射 CKO 回复到 UI → 经 _on_agent_responded → agent_response
        self.agent_responded.emit("CKO", response)

        # 记录到上下文
        self.ctx.add_message("CKO", response)

        # 通知 grounding 阶段完成
        self.debate_finished.emit(response)

    # ------------------------------------------------------------------
    #  Phase 2: 增强多轮辩论
    # ------------------------------------------------------------------
    def _run_debate(self):
        """增强多轮辩论 — 个性化风格 + 结构化投票 + 语义路由 + 历史引用。"""
        print(f"\n{'='*60}")
        print(f"[DEBATE-V2] 辩论阶段启动 | 轮次: {DEBATE_MIN_ROUNDS}-{DEBATE_MAX_ROUNDS}")
        print(f"{'='*60}\n")

        for round_num in range(1, DEBATE_MAX_ROUNDS + 1):
            if self._stop_flag:
                break

            print(f"\n[DEBATE-V2] ===== 第 {round_num} 轮辩论 =====\n")
            self.round_completed.emit(round_num)

            # [FIX] 竞态条件：在每轮开始前显式检查用户干预
            # 确保即使在 PM 发言前插入的指令也能被立即捕获
            last_msg = self.ctx.get_last_message()
            if last_msg and last_msg.role == "Commander" and last_msg.priority >= 10:
                print(f"[DEBATE-V2] 🔴 检测到用户干预 (Priority {last_msg.priority})，本轮交由 PM 优先处理。")
                # 不做额外操作，PM 的 Prompt 会自然包含此消息（如果 get_recent_history 包含它）
                # 但为了保险，我们可以短暂 sleep 确保 context 同步（虽然 list 是原子操作）
                time.sleep(0.1)

            # ---- PM 主持本轮 ----
            PersonalityEngine.thinking_delay()
            pm_prompt = self._build_debate_prompt("PM", round_num,
                                                  is_moderator=True)
            pm_agent = self.agents["PM"]
            pm_response = self._execute_single_agent(pm_agent, pm_prompt)
            if pm_response:
                self.ctx.add_message("PM", pm_response, round_num)
                self.agent_responded.emit("PM", pm_response)

            if self._stop_flag:
                break

            # ---- 路由器决定发言顺序 ----
            speakers = self.router.get_debate_order(round_num)
            previous_speaker = "PM"  # PM 是本轮第一个发言的
            previous_content = pm_response or ""

            for speaker_role in speakers:
                if self._stop_flag:
                    break

                # [NEW] 实时用户干预检测
                last_msg = self.ctx.history[-1] if self.ctx.history else None
                if last_msg and last_msg.role == "Commander":
                    print(f"[DEBATE-V2] 🔴 检测到用户干预，中断当前发言队列，转交 PM 处理。")
                    break

                PersonalityEngine.thinking_delay()
                agent = self._get_agent_by_role(speaker_role)
                if not agent:
                    continue

                # 构建增强 Prompt: 个性化 + 历史引用 + 立场态势 + 明确回应对象
                prompt = self._build_debate_prompt(
                    speaker_role, round_num,
                    previous_speaker=previous_speaker,
                    previous_content=previous_content
                )
                response = self._execute_single_agent(agent, prompt)
                if not response:
                    continue

                msg = self.ctx.add_message(speaker_role, response, round_num)
                self.agent_responded.emit(speaker_role, response)

                # 记录当前发言者，供下一位参考
                previous_speaker = speaker_role
                previous_content = response

                # ---- @提及处理 (被点名的角色回复) ----
                mentioned_roles = set()  # 记录已通过@提及回复的角色
                if msg.mentions:
                    for mentioned in msg.mentions:
                        target = self.router.get_response_to_mention([mentioned])
                        if target and target != speaker_role:
                            mentioned_roles.add(target)
                            print(f"[DEBATE-V2] {speaker_role} @{target} → 触发定向回复")
                            PersonalityEngine.thinking_delay()
                            reply_agent = self._get_agent_by_role(target)
                            if reply_agent:
                                reply_prompt = self._build_mention_prompt(
                                    target, speaker_role, response, round_num)
                                reply = self._execute_single_agent(
                                    reply_agent, reply_prompt)
                                if reply:
                                    self.ctx.add_message(target, reply, round_num)
                                    self.agent_responded.emit(target, reply)
                                    # 更新上一位发言者
                                    previous_speaker = target
                                    previous_content = reply

                # ---- 语义相关性触发 (主动响应) ----
                # 排除: 已在发言队列中的 / 自己 / PM / 已通过@提及回复的角色
                for other_role in self.router.ALL_ROLES:
                    if other_role in speakers or other_role == speaker_role:
                        continue
                    if other_role == "PM":
                        continue
                    if other_role in mentioned_roles:  # 已通过@提及回复，不重复触发
                        continue
                    if self.router.should_agent_respond(other_role, msg):
                        print(f"[DEBATE-V2] 语义路由: {other_role} 对 {speaker_role} 的消息感兴趣")
                        other_agent = self._get_agent_by_role(other_role)
                        if other_agent:
                            PersonalityEngine.thinking_delay()
                            reactive_prompt = self._build_reactive_prompt(
                                other_role, speaker_role, response, round_num)
                            reactive_reply = self._execute_single_agent(
                                other_agent, reactive_prompt)
                            if reactive_reply:
                                self.ctx.add_message(other_role, reactive_reply,
                                                     round_num)
                                self.agent_responded.emit(other_role,
                                                          reactive_reply)

                # ---- 冲突检测 ----
                if self.router.detect_conflict():
                    print(f"[DEBATE-V2] ⚠️ 冲突检测!")
                    if self.router.should_trigger_sub_debate():
                        self._run_sub_debate(round_num)

            # ---- 结构化投票 (每轮结束) ----
            if round_num >= DEBATE_MIN_ROUNDS:
                self._execute_structured_vote(round_num)
                if self.ctx.check_consensus():
                    print(f"\n[DEBATE-V2] ✅ 共识达成! 第 {round_num} 轮")
                    break
            print(f"[DEBATE-V2] 共识状态: {status}")

        # ---- PM 最终裁决 (含死锁熔断) ----
        final_prompt = ""
        if not self.ctx.check_consensus():
            print("\n[DEBATE-V2] ⚠️ 达到最大轮次仍未共识，触发 PM 独裁模式")
            final_prompt = (
                f"你是 PM。团队经过 {DEBATE_MAX_ROUNDS} 轮激烈辩论，仍无法达成一致。\n"
                f"{self.ctx.get_stance_summary()}\n\n"
                f"## 你的任务 (Dictator Mode)\n"
                f"- 作为最高决策者，请**强制指定**一个最终方案。\n"
                f"- 忽略少数派的纠结，基于 ROI 和可行性做决断。\n"
                f"- 必须以 '最终裁决：由于时间紧迫，我决定...' 开头。\n"
                f"- 简体中文。"
            )
        else:
            final_prompt = self._build_final_decision_prompt()

        final_resp = self._execute_single_agent(self.agents["PM"], final_prompt)
        if final_resp:
            self.ctx.add_message("PM", final_resp)
            self.agent_responded.emit("PM", final_resp)
        self.debate_finished.emit(self.ctx.get_stance_summary())

    # ------------------------------------------------------------------
    #  结构化投票
    # ------------------------------------------------------------------
    def _execute_structured_vote(self, round_num: int):
        """让每个辩论参与者投票：agree / disagree / conditional。"""
        print(f"[VOTE] 第 {round_num} 轮投票开始")
        voters = ["Arch", "Designer", "Coder"]
        for voter_role in voters:
            if self._stop_flag:
                break
            agent = self._get_agent_by_role(voter_role)
            if not agent:
                continue
            vote_prompt = (
                f"你是 {voter_role}，当前辩论到了投票环节。\n\n"
                f"{self.ctx.get_stance_summary()}\n\n"
                f"## 最近讨论\n{self.ctx.get_recent_history(limit=4)}\n\n"
                f"请用以下格式投票（必须严格遵守）：\n"
                f"VOTE: [agree/disagree/conditional]\n"
                f"REASON: [一句话理由]\n\n"
                f"简体中文回复，仅输出投票格式，不要其他内容。"
            )
            vote_resp = self._execute_single_agent(agent, vote_prompt)
            if vote_resp:
                # 解析投票
                stance = "neutral"
                reason = vote_resp
                if "agree" in vote_resp.lower():
                    stance = "agree"
                elif "disagree" in vote_resp.lower():
                    stance = "disagree"
                elif "conditional" in vote_resp.lower():
                    stance = "conditional"
                reason_match = re.search(r"REASON:\s*(.+)", vote_resp)
                if reason_match:
                    reason = reason_match.group(1).strip()
                self.ctx.register_vote(voter_role, stance, reason, round_num)
                self.agent_responded.emit(
                    voter_role, f"🗳️ [{stance.upper()}] {reason}")
                print(f"[VOTE] {voter_role}: {stance}")

    # ------------------------------------------------------------------
    #  子辩论 — 冲突解决
    # ------------------------------------------------------------------
    def _run_sub_debate(self, parent_round: int):
        """针对特定冲突进行子辩论。"""
        print(f"\n[SUB-DEBATE] 🔥 子辩论启动\n")
        disagrees = [r for r, s in self.ctx.consensus_tracker.items()
                     if s == "disagree"]
        if not disagrees:
            return
        for dissenter in disagrees[:2]:
            if self._stop_flag:
                break
            agent = self._get_agent_by_role(dissenter)
            if not agent:
                continue
            prompt = (
                f"你是 {dissenter}，你之前对方案提出了质疑。\n"
                f"现在请详细阐述你的反对理由，并给出替代方案。\n"
                f"请具体、有建设性、用数据或案例支撑。\n\n"
                f"## 上下文\n{self.ctx.get_recent_history(limit=5)}\n\n"
                f"简体中文回应："
            )
            resp = self._execute_single_agent(agent, prompt)
            if resp:
                self.ctx.add_message(dissenter, resp, parent_round)
                self.agent_responded.emit(dissenter, resp)

        # PM 调停
        med_prompt = (
            f"你是 PM，团队出现分歧。请综合各方意见给出折中方案。\n\n"
            f"## 上下文\n{self.ctx.get_recent_history(limit=8)}\n\n"
            f"简体中文裁决："
        )
        med_resp = self._execute_single_agent(self.agents["PM"], med_prompt)
        if med_resp:
            self.ctx.add_message("PM", med_resp, parent_round)
            self.agent_responded.emit("PM", med_resp)
        print(f"[SUB-DEBATE] 结束\n")

    # ------------------------------------------------------------------
    #  Phase 3: Production
    # ------------------------------------------------------------------
    def _run_production(self):
        """Coder 编码 + Tester 验证循环。"""
        print(f"\n[PRODUCTION] 生产阶段 | 最大循环: {PRODUCTION_MAX_LOOPS}\n")
        for loop in range(1, PRODUCTION_MAX_LOOPS + 1):
            if self._stop_flag:
                break
            print(f"\n[PRODUCTION] --- 循环 {loop} ---\n")
            
            # [FIX] Coder 工具动态挂载
            # 确保 Coder 拥有写文件的能力
            if not getattr(coder_agent, 'tools', None):
                try:
                    from tools.crew_tools import CodeWriterTool
                    coder_agent.tools = [CodeWriterTool()]
                    print(f"[PRODUCTION] ✅ 已为 Coder 挂载 CodeWriterTool")
                except Exception as e:
                    print(f"[PRODUCTION] ⚠️ 无法挂载 CodeWriterTool: {e}")

            # Coder
            coder_agent = self.agents["Coder"]
            coder_prompt = self._build_production_prompt("Coder", loop)
            coder_resp = self._execute_single_agent(coder_agent, coder_prompt)
            if coder_resp:
                self.ctx.add_message("Coder", coder_resp)
                self.agent_responded.emit("Coder", coder_resp)
            else:
                # [FIX] 断路器：Coder 失败则阻断后续流程
                print(f"[PRODUCTION] ❌ Coder 产出失败，跳过测试")
                self.agent_responded.emit("System", "⚠️ Coder 生成代码失败，跳过本轮测试。")
                continue

            if self._stop_flag:
                break
            PersonalityEngine.thinking_delay()
            # Tester
            tester_agent = self.agents["Tester"]
            tester_prompt = self._build_production_prompt("Tester", loop)
            tester_resp = self._execute_single_agent(tester_agent, tester_prompt)
            if tester_resp:
                msg = self.ctx.add_message("Tester", tester_resp)
                self.agent_responded.emit("Tester", tester_resp)
                if msg.sentiment == "positive" or "通过" in tester_resp:
                    print(f"[PRODUCTION] ✅ Tester 通过!")
                    break
                else:
                    print(f"[PRODUCTION] ❌ 需修复")
        self.production_finished.emit("生产阶段完成")

    # ------------------------------------------------------------------
    #  Phase 4: Review
    # ------------------------------------------------------------------
    def _run_review(self):
        """PM 最终审批。"""
        prompt = self._build_review_prompt()
        resp = self._execute_single_agent(self.agents["PM"], prompt)
        if resp:
            self.ctx.add_message("PM", resp)
            self.agent_responded.emit("PM", resp)
        self.debate_finished.emit("最终审批完成")

    # ------------------------------------------------------------------
    #  Phase X: Cloning (文档克隆)
    # ------------------------------------------------------------------
    def _run_cloning(self):
        """执行文档克隆任务。"""
        print(f"\n[CLONING] 文档克隆模式启动\n")
        self.agent_responded.emit("System", "🚀 进入文档克隆/重组模式...")
        
        # 尝试从历史中提取原型和新内容指令
        # 这里我们假设 CoderAgent 的 run_document_cloning 只需要 user_input 即可
        # 因为 CoderAgent 内部 logic 会再次解析，或者我们可以手动构造
        
        coder = self.agents["Coder"]
        user_input = self.user_input or (self.ctx.history[-1].content if self.ctx.history else "")
        
        # 为了兼容 Coder.run_document_cloning 的签名 (prototype, new_content)，
        # 我们这里简化处理：将整个 input 传给 LLM 让它自己去提取，或者
        # 我们显式要求用户提供。
        # 鉴于 CoderAgent._call_api 已经有了智能识别逻辑，我们这里直接调用 api
        # 或者 显式调用 run_document_cloning 如果我们在 Coder 中暴露了它。
        
        # Coder.run_document_cloning 是个方法，接受两个参数。
        # 我们尝试从上下文智能提取，如果提取不到，就让 Coder 自己处理。
        
        if hasattr(coder, "run_document_cloning"):
             # 尝试正则提取
            import re
            proto_match = re.search(r"【原型文档】[:\s]*([\s\S]*?)(?=【新内容】|$)", user_input)
            new_match = re.search(r"【新内容】[:\s]*([\s\S]*)", user_input)
            
            if proto_match and new_match:
                proto = proto_match.group(1).strip()
                new_c = new_match.group(1).strip()
                self.agent_responded.emit("System", f"📄 识别到原型文档 ({len(proto)} chars) 和新内容 ({len(new_c)} chars)")
                res = coder.run_document_cloning(proto, new_c)
                self.agent_responded.emit("Coder", res)
                self.production_finished.emit("克隆完成")
                return

        # Fallback: 让 Coder 作为普通 Agent 运行，但 Prompt 强调克隆
        prompt = (
            f"用户指令: {user_input}\n\n"
            f"任务类型: 文档克隆 (DOCUMENT CLONING)\n"
            f"请执行文档克隆任务。如果缺少信息，请询问用户。"
            f"如果已包含【原型文档】和【新内容】，请生成新的文档。"
        )
        resp = self._execute_single_agent(coder, prompt)
        if resp:
            self.ctx.add_message("Coder", resp)
            self.agent_responded.emit("Coder", resp)
        self.production_finished.emit("克隆任务结束")

    # ------------------------------------------------------------------
    #  Agent 调用
    # ------------------------------------------------------------------
    def _execute_single_agent(self, agent: Agent, prompt: str) -> Optional[str]:
        """使用 CrewAI 执行单个 Agent 任务。"""
        role_name = getattr(agent, 'role', 'Unknown')
        print(f"[EXEC] 正在执行 {role_name}...")
        try:
            temp_task = Task(description=prompt,
                             expected_output="简体中文的详细回复",
                             agent=agent)
            mini_crew = Crew(agents=[agent], tasks=[temp_task],
                             process=Process.sequential, verbose=True)
            result = mini_crew.kickoff()
            response = self._clean_response(str(result).strip())
            if response and len(response) > 3:
                print(f"[EXEC] {role_name} 成功 ({len(response)} 字)")
                return response
            else:
                print(f"[EXEC] {role_name} 返回空或过短响应")
                return None
        except Exception as e:
            error_msg = f"⚠️ {role_name} 执行失败: {str(e)[:200]}"
            print(f"[ERROR] {error_msg}")
            # 将错误信息也发送到 UI，避免静默失败
            self.agent_responded.emit(role_name.split("·")[0].strip(),
                                       error_msg)
            return None

    def _clean_response(self, text: str) -> str:
        """清理 CrewAI 输出噪声。"""
        # [FIX] 保护代码块：如果包含 ``` 则跳过正则清洗，防止误删代码
        if "```" in text:
            return text.strip()

        noise_patterns = [
            r'^Agent:.*$', r'^Task Completed.*$', r'^Name:.*$',
            r'^User Request:.*$', r'^Final Answer:?\s*',
        ]
        lines = text.split('\n')
        cleaned = [l for l in lines
                   if not any(re.match(p, l.strip(), re.IGNORECASE)
                              for p in noise_patterns)]
        return '\n'.join(cleaned).strip()

    # ------------------------------------------------------------------
    #  增强 Prompt 构建
    # ------------------------------------------------------------------
    def _build_debate_prompt(self, role: str, round_num: int,
                             is_moderator: bool = False,
                             previous_speaker: str = "",
                             previous_content: str = "") -> str:
        """构建增强辩论 Prompt — 三明治结构。

        [V2] 利用 LLM 的 Primacy + Recency Bias：
          开头：constraints_block（⛔ 死命令，结构化约束，短且醒目）
          中间：recent 讨论 + stance 态势 + style 风格
          末尾：mission_anchor + task 指令（放最后 = 最高注意力）

        Args:
            role: 当前发言角色
            round_num: 当前轮次
            is_moderator: 是否为 PM 主持人模式
            previous_speaker: 上一位发言者角色名
            previous_content: 上一位发言者的内容（截取前200字）
        """
        stance = self.ctx.get_stance_summary()
        recent = self.ctx.get_recent_history(limit=6)
        relevant = self.ctx.get_relevant_history(role, limit=3)
        style = PersonalityEngine.get_style_directive(role)
        mission_anchor = self.ctx.get_mission_anchor()

        # ---- [BLOCK 1] ⛔ 死命令约束区块（放最前 = Primacy Bias） ----
        constraints_block = ""
        if self.ctx.constraints:
            constraint_lines = "\n".join(
                f"  {i}. {c}" for i, c in enumerate(self.ctx.constraints, 1)
            )
            constraints_block = (
                f"⛔ 以下为本次任务的死命令，任何讨论不得违背：\n"
                f"{constraint_lines}\n\n"
            )

        # 根据任务类型注入约束
        task_type = self.ctx.task_type
        task_guidance = ""
        if task_type == "RESEARCH":
            task_guidance = (
                "【重要约束】当前任务为『调研/文档』。\n"
                "- 请聚焦于内容深度、逻辑结构和资料依据。\n"
                "- **严禁讨论具体的代码实现细节**，除非涉及技术原理说明。\n"
            )
        elif task_type == "DESIGN":
            task_guidance = (
                "【重要约束】当前任务为『设计/方案』。\n"
                "- 请聚焦于用户体验、视觉风格和交互流程。\n"
                "- **严禁讨论后端代码实现**，关注界面表现层。\n"
            )
        else:  # ENGINEERING
            task_guidance = (
                "【重要约束】当前任务为『工程实现』。\n"
                "- 请聚焦于技术架构、代码规范和可行性。\n"
                "- 鼓励讨论具体实现细节。\n"
            )

        # [Role-Specific Guidance]
        # 针对 Coder 在非代码任务中的特殊引导
        role_guidance = ""
        if role == "Coder":
            if task_type == "DESIGN":
                role_guidance = (
                    "\n【Coder 特别指令】当前是『设计阶段』，请扮演【技术顾问】。\n"
                    "- 不要写代码！\n"
                    "- 你的职责是评估设计方案的【工程可行性】和【成本】。\n"
                    "- 指出看似简单但实现极难的交互（'这个动画会卡顿'）。\n"
                    "- 建议复用现有组件库以节省工期。\n"
                )
            elif task_type == "RESEARCH":
                role_guidance = (
                    "\n【Coder 特别指令】当前是『调研阶段』，请扮演【数据技术专家】。\n"
                    "- 不要写代码！\n"
                    "- 你的职责是评估【数据获取难度】和【API 限制】。\n"
                    "- 为产品构想提供技术边界支持（'这个功能很难实现，因为没有开放API'）。\n"
                )

        # ================================================================
        #  三明治结构：[死命令] → [中间内容] → [锚点 + 任务指令]
        # ================================================================

        if is_moderator:
            # ---- PM 主持人模式（三明治结构） ----
            return (
                # --- 🔝 顶部：死命令约束 (Primacy Bias) ---
                f"{constraints_block}"
                f"{task_guidance}\n"
                f"{role_guidance}"
                f"你是 PM（项目经理），正在主持第 {round_num} 轮方案讨论。\n\n"
                # --- 🔄 中间：讨论上下文 + 态势 + 风格 ---
                f"## 最近讨论\n{recent}\n\n"
                f"[风格指令] {style}\n\n"
                f"{stance}\n\n"
                # --- 🔚 末尾：锚点 + 任务指令 (Recency Bias = 最高注意力) ---
                f"{mission_anchor}\n\n"
                f"## 你的任务（最终命令）\n"
                f"- 必须基于【原始需求锚点】来评估当前的讨论方向，防止跑题。\n"
                f"{'- 第一轮，请基于任务协议提出初始方案框架，@Arch 和 @Designer 发表意见。' if round_num == 1 else '- 总结上轮进展和分歧，引导团队聚焦核心问题。'}\n"
                f"- 如团队即将共识，说 '我同意现有方案' 或 '方案通过'\n"
                f"- 有分歧时 @相关角色\n"
                f"- 简体中文，300字内"
            )

        # ---- 普通角色 Prompt（三明治结构） ----
        # 明确指定回应对象
        reply_directive = ""
        if previous_speaker and previous_content:
            reply_directive = f"请回应 {previous_speaker} 的观点：\"{previous_content[:150]}...\"\n\n"

        # 相关历史引用
        relevant_section = ""
        if relevant:
            relevant_section = f"{relevant}\n\n"

        return (
            # --- 🔝 顶部：死命令约束 (Primacy Bias) ---
            f"{constraints_block}"
            f"{task_guidance}\n"
            f"{role_guidance}"
            f"你是 {role}，正在 War Room 第 {round_num} 轮辩论。\n\n"
            # --- 🔄 中间：讨论上下文 + 态势 + 风格 ---
            f"{reply_directive}"
            f"## 最近讨论\n{recent}\n\n"
            f"{relevant_section}"
            f"[风格指令] {style}\n\n"
            f"{stance}\n\n"
            # --- 🔚 末尾：锚点 + 任务指令 (Recency Bias = 最高注意力) ---
            f"{mission_anchor}\n\n"
            f"## 你的任务（最终命令）\n"
            f"- 必须基于【原始需求锚点】来评估当前的讨论方向。\n"
            f"- 如果发现讨论偏题，请立即拉回。\n"
            f"- 直接回应 {previous_speaker if previous_speaker else '上一位发言者'}，"
            f"必须明确表态（同意/反对/有条件同意）\n"
            f"- 引用之前的讨论: '正如 @某某 在之前提到的...'\n"
            f"- 如需信息请 @角色名 点名提问\n"
            f"- 如反对请 @角色名 指出分歧并提替代方案\n"
            f"- 简体中文，300字内"
        )

    def _build_mention_prompt(self, responder: str, mentioner: str,
                              content: str, round_num: int) -> str:
        """被 @提及时的回应 Prompt。"""
        style = PersonalityEngine.get_style_directive(responder)
        return (
            f"你是 {responder}。{mentioner} 在讨论中 @提及了你：\n"
            f"「{content[:300]}」\n\n"
            f"[风格指令] {style}\n\n"
            f"## 上下文\n{self.ctx.get_recent_history(limit=4)}\n\n"
            f"直接回应 {mentioner}，明确同意或反对。简体中文，200字内。"
        )

    def _build_reactive_prompt(self, responder: str, trigger_role: str,
                               trigger_content: str, round_num: int) -> str:
        """语义路由触发的主动回应 Prompt。"""
        style = PersonalityEngine.get_style_directive(responder)
        return (
            f"你是 {responder}。{trigger_role} 刚说了一些与你专业相关的内容：\n"
            f"「{trigger_content[:200]}」\n\n"
            f"[风格指令] {style}\n\n"
            f"你觉得有必要从 {responder} 的专业角度补充或质疑。\n"
            f"请主动插话，用 @{trigger_role} 开头回应。简体中文，200字内。"
        )

    def _build_final_decision_prompt(self) -> str:
        """PM 最终裁决 Prompt。"""
        return (
            f"你是 PM，辩论即将结束。\n\n"
            f"{self.ctx.get_stance_summary()}\n\n"
            f"## 完整讨论\n{self.ctx.get_recent_history(limit=15)}\n\n"
            f"## 投票记录\n" +
            self._format_voting_results() +  # [FIX] 投票去重
            f"\n\n请做最终决定，以 '方案通过：...' 开头。简体中文。"
        )

    def _format_voting_results(self) -> str:
        """格式化投票结果，自动去重（仅保留每人最新的一票）。"""
        latest_votes = {}
        # 顺序遍历，后面的覆盖前面的
        for v in self.ctx.vote_records:
            latest_votes[v['role']] = v
        
        lines = []
        for role, v in latest_votes.items():
            lines.append(f"- {role}: {v['stance']} ({v['reason']})")
        return "\n".join(lines)

    def _build_production_prompt(self, role: str, loop: int) -> str:
        """生产阶段 Prompt - 动态适配任务类型"""
        recent = self.ctx.get_recent_history(limit=6)
        task_type = self.ctx.task_type  # 获取任务类型

        extra = ""
        if role == "Coder":
            if task_type == "RESEARCH":
                # 针对调研任务的指令
                extra = (
                    f"{'根据最终方案撰写深度研究报告。' if loop == 1 else '根据反馈优化报告内容。'}\n"
                    f"- 必须输出为标准的 Markdown 格式\n"
                    f"- 包含：核心结论、理论框架、数据支持、潜在风险\n"
                    f"- 严禁输出 Python 代码，除非是用于数据可视化的脚本"
                )
            elif task_type == "DESIGN":
                # 针对设计任务的指令
                extra = (
                    f"{'根据最终方案输出详细设计文档。' if loop == 1 else '根据反馈完善设计细节。'}\n"
                    f"- 输出完整的 Markdown 设计规范\n"
                    f"- 包含：Mermaid 流程图、界面描述、交互逻辑\n"
                    f"- 不需要实现具体业务代码"
                )
            else: # ENGINEERING / DEFAULT
                # 原有的代码开发指令
                extra = (
                    f"{'根据最终方案编写代码。' if loop == 1 else '根据 Tester 反馈修复。'}\n"
                    f"- 遵循 Google 编程规范，详尽中文注释\n"
                    f"- 必须使用 `code_writer` 工具将代码保存到文件（如 `code/xxx.py`）\n"
                    f"- 同时在回答中展示核心代码片段"
                )
        else:
            # Tester also needs to adapt
            if task_type == "RESEARCH":
                extra = (
                    f"评审 Coder 的研究报告。\n"
                    f"- 检查逻辑漏洞、数据来源、论证深度\n"
                    f"- 通过说 '评审通过 ✅'，否则提出修改意见"
                )
            elif task_type == "DESIGN":
                extra = (
                    f"评审 Coder 的设计文档。\n"
                    f"- 检查交互逻辑完整性、视觉规范一致性\n"
                    f"- 通过说 '评审通过 ✅'，否则提出修改意见"
                )
            else: # ENGINEERING
                extra = (
                    f"验证 Coder 产出。\n"
                    f"- 检查逻辑、错误处理、规范性\n"
                    f"- 通过说 '测试通过 ✅'，否则说明问题和修复建议"
                )

        return (
            f"你是 {role}，生产阶段第 {loop} 轮。\n\n"
            f"## 上下文\n{recent}\n\n## 任务\n{extra}\n- 简体中文"
        )

    def _build_review_prompt(self) -> str:
        """PM 最终审批 Prompt。"""
        return (
            f"你是 PM，请对整个项目进行最终审批。\n\n"
            f"## 回顾\n{self.ctx.get_recent_history(limit=20)}\n\n"
            f"## 任务\n- 审查所有产出\n- 判断是否满足需求\n"
            f"- 给出审批意见\n- 简体中文"
        )

    def _get_agent_by_role(self, role: str) -> Optional[Agent]:
        """根据角色名获取 Agent。"""
        # [MODIFIED] 从 self.agents 字典获取
        return self.agents.get(role)


# ==============================================================================
#  7. ProxyAgent — UI 打字状态兼容层
# ==============================================================================

class ProxyAgent(QObject):
    """代理对象，模拟 Agent 的打字状态信号。

    供 main.py 连接 typing_started / typing_finished 信号。
    """
    typing_started = pyqtSignal()
    typing_finished = pyqtSignal()

    def __init__(self, role: str, parent=None):
        super().__init__(parent)
        self.role = role


# ==============================================================================
#  8. StateController — 状态机控制
# ==============================================================================

class StateController(QObject):
    """简单状态机，管理应用阶段转换。"""
    state_changed = pyqtSignal(str, str)  # state_id, description

    STATE_DESC = {
        AppState.IDLE: "空闲状态",
        AppState.GROUNDING: "CKO 需求打磨中...",
        AppState.DEBATE: "方案博弈中...",
        AppState.PRODUCTION: "编码阶段...",
        AppState.VERIFICATION: "验证中...",
        AppState.DELIVERY: "交付汇总中...",
        AppState.COMPLETED: "任务完成",
    }

    def __init__(self, parent=None):
        super().__init__(parent)
        self.current_state = AppState.IDLE

    def transition(self, new_state: str):
        """切换状态并发射信号。"""
        self.current_state = new_state
        desc = self.STATE_DESC.get(new_state, new_state)
        print(f"[STATE] {self.current_state} → {new_state}: {desc}")
        self.state_changed.emit(new_state, desc)


# ==============================================================================
#  9. CommanderCrew — 主编排器 (替代 CrewManager)
# ==============================================================================

class CommanderCrew(QObject):
    """增强型多 Agent 编排器 — 信号接口与 CrewManager 完全兼容。

    信号:
      agent_response(role, content)   — 分发到 main.py 的 UI
      state_changed(state_id, desc)   — 状态变更通知
      error_occurred(msg)             — 错误通知
      workflow_completed()            — 全部完成
      debate_round_info(int)          — 辩论轮次信息 (可选)
    """
    agent_response = pyqtSignal(str, str)
    state_changed = pyqtSignal(str, str)
    error_occurred = pyqtSignal(str)
    workflow_completed = pyqtSignal()
    debate_round_info = pyqtSignal(int)

    def __init__(self, parent=None):
        super().__init__(parent)

        # 状态机
        self.state_ctrl = StateController(self)
        self.state_ctrl.state_changed.connect(self.state_changed)

        # 上下文 & 路由器
        self.ctx = WarRoomContext()
        self.router = DebateRouter(self.ctx)

        # 当前 Worker
        self._current_worker: Optional[DebateWorker] = None
        self._user_input: str = ""

        # ProxyAgent — 供 main.py 连接 typing 信号
        self.cko = ProxyAgent("CKO", self)
        self.pm = ProxyAgent("PM", self)
        self.arch = ProxyAgent("Arch", self)
        self.designer = ProxyAgent("Designer", self)
        self.coder = ProxyAgent("Coder", self)
        self.tester = ProxyAgent("Tester", self)

        # 角色 → ProxyAgent 映射
        self._proxy_map = {
            "CKO": self.cko, "PM": self.pm, "Arch": self.arch,
            "Designer": self.designer, "Coder": self.coder,
            "Tester": self.tester,
        }
        
        # [NEW] 保持 agents 引用以便非 Worker 线程也能访问（如 handle_user_intervention）
        # 初始化默认 agents (无回调)
        self.default_agents = create_agents() 

    # ==================================================================
    #  公开接口 — 与 CrewManager 兼容
    # ==================================================================

    def start_mission(self, user_input: str):
        """Phase 1: CKO Grounding.

        由 Bridge 面板的 message_sent 信号触发。
        """
        self._user_input = user_input
        self.state_ctrl.transition(AppState.GROUNDING)
        self._launch_worker("grounding", user_input)

    def confirm_project(self):
        """Phase 2: 启动辩论 (或克隆任务)。

        由 main.py 的 Confirm Project 按钮触发。
        """
        if self.ctx.task_type == "CLONING":
            # [NEW] 克隆模式分支
            # 我们暂时复用 PRODUCTION 状态，或者定义新的 CLONING 状态
            # 这里简单复用 PRODUCTION，但 worker mode 为 'cloning'
            self.state_ctrl.transition(AppState.PRODUCTION)
            self._launch_worker("cloning")
        else:
            self.state_ctrl.transition(AppState.DEBATE)
            self._launch_worker("debate")

    def handle_user_intervention(self, message: str):
        """用户干预 — 高优先级广播。

        将用户消息注入上下文，PM 立即响应。
        """
        print(f"[INTERVENTION] 用户干预: {message[:50]}...")

        # 1. 高优先级注入上下文
        self.ctx.add_message("Commander", message, priority=10)
        self.agent_response.emit("Commander", message)

        # 2. 如果没有正在运行的 Worker，创建一个临时 PM 响应
        if self._current_worker is None or not self._current_worker.isRunning():
            self._launch_intervention_response(message)

    def new_session(self):
        """重置会话。"""
        self.stop_all()
        self.ctx.reset()
        self._user_input = ""
        self.state_ctrl.transition(AppState.IDLE)

    def stop_all(self):
        """停止所有正在运行的 Worker。"""
        if self._current_worker and self._current_worker.isRunning():
            self._current_worker.stop()
            self._current_worker.quit()
            self._current_worker.wait(3000)
            self._current_worker = None

    # ==================================================================
    #  Worker 生命周期管理
    # ==================================================================

    def _launch_worker(self, mode: str, user_input: str = ""):
        """创建并启动 DebateWorker。"""
        if self._current_worker and self._current_worker.isRunning():
            self._current_worker.stop()
            self._current_worker.quit()
            self._current_worker.wait(2000)

        worker = DebateWorker(mode, self.ctx, self.router,
                              user_input or self._user_input, self)
        # 连接信号
        worker.agent_responded.connect(self._on_agent_responded)
        worker.round_completed.connect(self.debate_round_info)
        worker.debate_finished.connect(self._on_debate_finished)
        worker.production_finished.connect(self._on_production_finished)
        worker.error_occurred.connect(self.error_occurred)

        self._current_worker = worker
        worker.start()

    def _launch_intervention_response(self, message: str):
        """针对用户干预，让 PM 立即响应。"""
        worker = DebateWorker("review", self.ctx, self.router,
                              message, self)
        worker.agent_responded.connect(self._on_agent_responded)
        worker.debate_finished.connect(lambda _: None)
        worker.error_occurred.connect(self.error_occurred)
        self._current_worker = worker
        worker.start()

    # ==================================================================
    #  信号处理
    # ==================================================================

    def _on_agent_responded(self, role: str, content: str):
        """转发 Agent 回复到 main.py，并触发打字状态。"""
        proxy = self._proxy_map.get(role)
        if proxy:
            proxy.typing_started.emit()

        # 发射主信号 → main.py 的 _dispatch_agent_response
        self.agent_response.emit(role, content)

        if proxy:
            proxy.typing_finished.emit()

    def _on_debate_finished(self, summary: str):
        """辩论结束 → 进入生产阶段。"""
        current = self.state_ctrl.current_state
        if current == AppState.GROUNDING:
            # Grounding 完成 → 等待确认
            self.state_ctrl.transition("AWAITING_CONFIRM")
            self.agent_response.emit("系统",
                                     "📋 CKO 需求分析完成，请确认项目进入辩论阶段。")
        elif current == AppState.DEBATE:
            # 辩论完成 → 生产
            self.agent_response.emit("系统",
                                     f"🏁 辩论结束\n{summary}")
            self.state_ctrl.transition(AppState.PRODUCTION)
            self._launch_worker("production")
        else:
            self.workflow_completed.emit()

    def _on_production_finished(self, result: str):
        """生产完成 → 最终审批。"""
        self.agent_response.emit("系统", "🔍 生产完成，进入最终审批...")
        self.state_ctrl.transition(AppState.VERIFICATION)
        self._launch_worker("review")

