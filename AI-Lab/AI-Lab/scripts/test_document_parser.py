"""
test_document_parser.py - 文档解析器综合测试
============================================
测试 document_parser.py 的全内容类型解析能力。

安全性审计:
  ✅ 测试文件由脚本自行创建和清理
  ✅ 不涉及外部 API 调用
"""

import sys
import os
import json
import tempfile
import shutil

# 确保项目根目录在 path 中
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from tools.document_parser import parse_document, ParsedDocument


def test_text_file():
    """测试纯文本解析。"""
    print("[1/5] 测试纯文本解析...")
    
    # 创建临时 txt 文件
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt",
                                      delete=False, encoding="utf-8") as f:
        f.write("这是第一行\n这是第二行\n第三行：包含中文和 English\n")
        tmp_path = f.name

    try:
        result = parse_document(tmp_path)
        assert result.format == "txt", f"格式错误: {result.format}"
        assert "这是第一行" in result.text, "文本未正确提取"
        assert result.metadata.get("编码") == "utf-8", f"编码检测错误: {result.metadata}"
        assert not result.errors, f"不应有错误: {result.errors}"
        
        prompt = result.to_prompt_text()
        assert "正文内容" in prompt, "Prompt 格式不正确"
        print(f"  ✅ 纯文本解析通过 (文本 {len(result.text)} 字符)")
    finally:
        os.unlink(tmp_path)
    return True


def test_markdown_file():
    """测试 Markdown 解析。"""
    print("[2/5] 测试 Markdown 解析...")
    
    md_content = """# 项目方案

## 1. 背景

这是一个 **测试项目**，用于验证文档解析功能。

## 2. 技术栈

- Python 3.10+
- PyQt6
- CrewAI

## 3. 代码示例

```python
def hello():
    print("Hello, World!")
```
"""
    
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md",
                                      delete=False, encoding="utf-8") as f:
        f.write(md_content)
        tmp_path = f.name

    try:
        result = parse_document(tmp_path)
        assert result.format == "md"
        assert "测试项目" in result.text
        assert "Python 3.10" in result.text
        print(f"  ✅ Markdown 解析通过 (文本 {len(result.text)} 字符)")
    finally:
        os.unlink(tmp_path)
    return True


def test_docx_full_content():
    """测试 Word 文档全内容解析 (文本/表格/元数据)。"""
    print("[3/5] 测试 Word 文档全内容解析...")
    
    try:
        import docx
        from docx.shared import Inches, Pt
    except ImportError:
        print("  ⚠️ python-docx 未安装，跳过 (pip install python-docx)")
        return True
    
    # 创建测试 Word 文档
    doc = docx.Document()
    
    # 元数据
    doc.core_properties.title = "测试文档标题"
    doc.core_properties.author = "AI-Lab 测试"
    doc.core_properties.subject = "全内容解析测试"
    
    # 标题和段落
    doc.add_heading("第一章 项目概述", level=1)
    doc.add_paragraph("这是项目概述的正文内容。包含多段中文文字。")
    doc.add_paragraph("第二段落：描述项目的技术架构和实现方案。")
    
    doc.add_heading("第二章 技术方案", level=2)
    doc.add_paragraph("采用 Python + PyQt6 技术栈进行开发。")
    
    # 表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    headers = ["组件", "技术", "状态"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["前端", "PyQt6", "完成"],
        ["后端", "CrewAI", "进行中"],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val
    
    # 页眉
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "AI-Lab-Commander 测试文档"
    
    # 页脚
    footer = section.footer
    footer.paragraphs[0].text = "机密文档 - 内部使用"
    
    # 保存
    tmp_path = os.path.join(tempfile.gettempdir(), "test_full_content.docx")
    doc.save(tmp_path)
    
    try:
        result = parse_document(tmp_path)
        
        # 验证格式
        assert result.format == "docx", f"格式错误: {result.format}"
        
        # 验证元数据
        assert result.metadata.get("标题") == "测试文档标题", f"标题错误: {result.metadata}"
        assert result.metadata.get("作者") == "AI-Lab 测试", f"作者错误: {result.metadata}"
        
        # 验证正文
        assert "项目概述" in result.text, "标题未提取"
        assert "技术架构" in result.text, "段落未提取"
        
        # 验证表格
        assert len(result.tables) >= 1, f"表格未提取: {len(result.tables)}"
        assert any("PyQt6" in str(row) for table in result.tables for row in table), "表格内容缺失"
        
        # 验证页眉
        assert any("测试文档" in h for h in result.headers), f"页眉未提取: {result.headers}"
        
        # 验证页脚
        assert any("机密文档" in f for f in result.footers), f"页脚未提取: {result.footers}"
        
        # 验证 Prompt 输出
        prompt = result.to_prompt_text()
        assert "文档信息" in prompt, "Prompt 缺少元数据"
        assert "正文内容" in prompt, "Prompt 缺少正文"
        assert "表格内容" in prompt, "Prompt 缺少表格"
        assert "页眉" in prompt, "Prompt 缺少页眉"
        assert "页脚" in prompt, "Prompt 缺少页脚"
        
        print(f"  ✅ Word 全内容解析通过:")
        print(f"     元数据: ✓ | 正文: ✓ | 表格: {len(result.tables)}个 ✓")
        print(f"     页眉: {len(result.headers)}个 ✓ | 页脚: {len(result.footers)}个 ✓")
        print(f"     图片: {len(result.images_base64)}张 | 批注: {len(result.comments)}条")
        print(f"     Prompt 输出: {len(prompt)} 字符")
    finally:
        os.unlink(tmp_path)
    return True


def test_unsupported_format():
    """测试不支持格式的优雅降级。"""
    print("[4/5] 测试不支持格式的降级...")
    
    with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
        f.write(b"binary data")
        tmp_path = f.name
    
    try:
        result = parse_document(tmp_path)
        assert len(result.errors) > 0, "应有错误信息"
        assert "不支持" in result.errors[0], f"错误类型不正确: {result.errors}"
        print(f"  ✅ 不支持格式降级通过: {result.errors[0]}")
    finally:
        os.unlink(tmp_path)
    return True


def test_nonexistent_file():
    """测试不存在文件的错误处理。"""
    print("[5/5] 测试不存在文件的错误处理...")
    
    result = parse_document("/nonexistent/file.docx")
    assert len(result.errors) > 0, "应有错误信息"
    assert "不存在" in result.errors[0], f"错误类型不正确: {result.errors}"
    print(f"  ✅ 文件不存在处理通过: {result.errors[0]}")
    return True


def main():
    """主测试入口"""
    print("=" * 60)
    print("  AI-Lab-Commander - 文档解析器综合测试")
    print("=" * 60)
    print()

    results = {}
    tests = [
        ("纯文本解析", test_text_file),
        ("Markdown解析", test_markdown_file),
        ("Word全内容解析", test_docx_full_content),
        ("不支持格式降级", test_unsupported_format),
        ("不存在文件处理", test_nonexistent_file),
    ]

    for name, test_fn in tests:
        try:
            results[name] = test_fn()
        except Exception as e:
            import traceback
            print(f"  ❌ {name}: {e}")
            traceback.print_exc()
            results[name] = False
        print()

    # 汇总
    print("=" * 60)
    print("  Summary Report")
    print("=" * 60)
    passed = sum(1 for v in results.values() if v)
    total = len(results)

    for name, ok in results.items():
        status = "✅" if ok else "❌"
        print(f"  {status} {name}")

    print(f"\n  Result: {passed}/{total} passed")
    print("=" * 60)
    return passed == total


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
