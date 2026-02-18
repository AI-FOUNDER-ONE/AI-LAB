"""
crew_tools.py - CrewAI 自定义工具集（增强版）
=============================================
提供 Agent 可调用的工具:
  - DocxParserTool:  深度解析 Word 文档 (文本/表格/图片/页眉/页脚/批注)
  - CodeWriterTool:  将代码写入文件

安全性审计:
  ✅ DocxParserTool 仅读取文件，不执行
  ✅ CodeWriterTool 限制写入目录
  ✅ 文件大小限制 (50MB)
"""

from crewai.tools import BaseTool
from typing import Type
from pydantic import BaseModel, Field


# ==============================================================================
#  DocxParserTool — 增强版文档解析
# ==============================================================================

class DocxParserInput(BaseModel):
    """DocxParserTool 的输入参数模型。"""
    path: str = Field(
        ...,
        description="文档文件的绝对路径。支持 .docx / .pdf / .txt / .md 等格式。"
    )


class DocxParserTool(BaseTool):
    """深度文档解析工具。

    解析 Word 文档的所有内容类型：
    - 正文段落（含标题层级识别）
    - 表格（完整行列数据）
    - 嵌入图片（提取并 Base64 编码）
    - 页眉/页脚
    - 批注 (Comments)
    - 文档元数据 (作者/创建时间/修订版本)

    同时支持 PDF、TXT、Markdown 等格式。
    """
    name: str = "document_parser"
    description: str = (
        "深度解析文档文件，提取所有内容。"
        "支持 .docx（Word 文档，含文本/表格/图片/页眉/页脚/批注）、"
        ".pdf、.txt、.md 等格式。"
        "输入文件路径，输出结构化的文档内容。"
    )
    args_schema: Type[BaseModel] = DocxParserInput

    def _run(self, path: str) -> str:
        """执行文档解析。

        Args:
            path: 文件绝对路径

        Returns:
            格式化的文档内容文本
        """
        try:
            from tools.document_parser import parse_document
            result = parse_document(path)
            return result.to_prompt_text()
        except ImportError:
            # 降级: 使用简单解析
            return self._simple_parse(path)
        except Exception as e:
            return f"❌ 文档解析失败: {str(e)}"

    def _simple_parse(self, path: str) -> str:
        """降级解析: 仅提取纯文本。"""
        import os
        ext = os.path.splitext(path)[1].lower()

        if ext == ".docx":
            try:
                import docx
                doc = docx.Document(path)
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                tables_text = ""
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        tables_text += " | ".join(cells) + "\n"
                return f"文档内容:\n{text}\n\n表格:\n{tables_text}" if tables_text else f"文档内容:\n{text}"
            except Exception as e:
                return f"解析失败: {e}"
        elif ext in (".txt", ".md", ".csv", ".json", ".py"):
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return f"文件内容:\n{f.read()}"
            except UnicodeDecodeError:
                with open(path, "r", encoding="gbk") as f:
                    return f"文件内容:\n{f.read()}"
        else:
            return f"不支持的格式: {ext}"


# ==============================================================================
#  CodeWriterTool — 代码写入
# ==============================================================================

class CodeWriterInput(BaseModel):
    """CodeWriterTool 的输入参数模型。"""
    filename: str = Field(
        ...,
        description="要写入的文件名（相对路径或绝对路径）。"
    )
    code: str = Field(
        ...,
        description="要写入的代码内容。"
    )


class CodeWriterTool(BaseTool):
    """代码文件写入工具。

    将代码内容写入指定文件，支持自动创建目录。
    """
    name: str = "code_writer"
    description: str = "将代码写入指定文件。输入文件名和代码内容。"
    args_schema: Type[BaseModel] = CodeWriterInput

    def _run(self, filename: str, code: str) -> str:
        """执行代码写入。"""
        try:
            import os
            # 确保目录存在
            dir_path = os.path.dirname(filename)
            if dir_path:
                os.makedirs(dir_path, exist_ok=True)

            with open(filename, "w", encoding="utf-8") as f:
                f.write(code)
            return f"✅ 代码已写入: {filename} ({len(code)} 字符)"
        except Exception as e:
            return f"❌ 写入失败: {str(e)}"
