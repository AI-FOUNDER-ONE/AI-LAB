"""
docx_generator_tool.py - DOCX 文档生成工具
==========================================
供 Coder（执行官）使用，将 Markdown 内容转换为结构化 Word 文档。
适用于项目计划书、可研报告、立项申请书等非代码交付物。

安全性审计:
  ✅ 文件写入限制在 data/generated_docs/ 目录
  ✅ 输入内容经过清洗处理
  ✅ 使用 python-docx 库，无外部进程调用
"""

import os
import re
import time
from typing import Type
from pydantic import BaseModel, Field
from tools.base_tool import BaseTool


class DocxGeneratorInput(BaseModel):
    """DocxGeneratorTool 的输入参数模型。"""
    content: str = Field(
        ...,
        description=(
            "Markdown 格式的文档内容。支持以下语法：\n"
            "- # 一级标题 / ## 二级标题 / ### 三级标题\n"
            "- **加粗文本** / *斜体文本*\n"
            "- - 无序列表项\n"
            "- 1. 有序列表项\n"
            "- ```代码块```\n"
            "- 普通段落文本"
        )
    )
    doc_title: str = Field(
        default="",
        description="文档标题（显示在封面页）。留空则从内容第一个标题提取。"
    )
    filename: str = Field(
        default="",
        description="输出文件名（不含扩展名）。留空则自动生成。"
    )


class DocxGeneratorTool(BaseTool):
    """DOCX 文档生成工具。

    将 Markdown 格式的内容转换为专业的 Word 文档。
    自动识别标题层级、列表、代码块等结构。

    使用场景:
      - 工程开发项目计划书（含 WBS 分解）
      - 科研项目立项申请书
      - 可行性研究报告
      - 工业设计方案文档
    """
    name: str = "docx_document_generator"
    description: str = (
        "将 Markdown 内容转换为 Word (.docx) 文档。"
        "输入 Markdown 格式的文本，输出保存的 .docx 文件路径。"
        "适用于生成项目计划书、可研报告、立项申请等正式文档。"
    )
    args_schema: Type[BaseModel] = DocxGeneratorInput

    def _run(self, content: str, doc_title: str = "", filename: str = "") -> str:
        """执行 DOCX 文档生成。

        Args:
            content: Markdown 格式的文档内容
            doc_title: 可选的文档标题
            filename: 可选的输出文件名

        Returns:
            生成文件的路径信息
        """
        # ---------- 1. 确定输出路径 ----------
        from config import GENERATED_DOCS_DIR

        if not filename:
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}"

        filename = filename.replace(os.sep, "_").replace("/", "_")
        output_path = os.path.join(GENERATED_DOCS_DIR, f"{filename}.docx")

        # ---------- 2. 尝试使用 python-docx 生成 ----------
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return "❌ python-docx 未安装。请运行: pip install python-docx"

        try:
            doc = Document()

            # ---------- 3. 添加标题页 ----------
            if not doc_title:
                # 从内容第一行提取标题
                first_line = content.strip().split("\n")[0]
                doc_title = re.sub(r'^#+\s*', '', first_line).strip() or "文档"

            title_para = doc.add_heading(doc_title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加生成时间
            time_para = doc.add_paragraph(
                f"生成时间: {time.strftime('%Y-%m-%d %H:%M:%S')}"
            )
            time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in time_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)

            doc.add_page_break()

            # ---------- 4. 解析 Markdown 内容 ----------
            lines = content.split("\n")
            in_code_block = False
            code_buffer = []

            for line in lines:
                stripped = line.strip()

                # 代码块处理
                if stripped.startswith("```"):
                    if in_code_block:
                        # 代码块结束 — 输出缓冲
                        code_text = "\n".join(code_buffer)
                        code_para = doc.add_paragraph()
                        code_run = code_para.add_run(code_text)
                        code_run.font.name = "Consolas"
                        code_run.font.size = Pt(9)
                        code_run.font.color.rgb = RGBColor(40, 40, 40)
                        code_buffer.clear()
                        in_code_block = False
                    else:
                        in_code_block = True
                    continue

                if in_code_block:
                    code_buffer.append(line)
                    continue

                # 空行跳过
                if not stripped:
                    continue

                # 标题 (# / ## / ###)
                heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
                if heading_match:
                    level = len(heading_match.group(1))
                    heading_text = heading_match.group(2)
                    doc.add_heading(heading_text, level=min(level, 4))
                    continue

                # 无序列表 (- 或 *)
                list_match = re.match(r'^[-*]\s+(.*)', stripped)
                if list_match:
                    doc.add_paragraph(
                        list_match.group(1), style='List Bullet'
                    )
                    continue

                # 有序列表 (1. 2. ...)
                olist_match = re.match(r'^\d+\.\s+(.*)', stripped)
                if olist_match:
                    doc.add_paragraph(
                        olist_match.group(1), style='List Number'
                    )
                    continue

                # 普通段落 — 处理加粗和斜体
                para = doc.add_paragraph()
                self._add_formatted_text(para, stripped)

            # ---------- 5. 保存文档 ----------
            doc.save(output_path)

            # 统计信息
            word_count = len(content.replace("\n", " ").split())
            return (
                f"✅ Word 文档已生成\n"
                f"  文件路径: {output_path}\n"
                f"  文档标题: {doc_title}\n"
                f"  约 {word_count} 字"
            )

        except Exception as e:
            return f"❌ 文档生成失败: {e}"

    @staticmethod
    def _add_formatted_text(paragraph, text: str):
        """解析 Markdown 行内格式并添加到段落。

        支持 **加粗** 和 *斜体* 语法。

        Args:
            paragraph: python-docx 段落对象
            text: 待处理的文本行
        """
        # 简单的行内格式解析：**bold** 和 *italic*
        pattern = r'(\*\*.*?\*\*|\*.*?\*|[^*]+)'
        parts = re.findall(pattern, text)

        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                # 加粗
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                # 斜体
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)
