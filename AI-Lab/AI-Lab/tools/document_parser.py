"""
document_parser.py - 统一文档解析器
=====================================
支持多种格式的文档解析，提取结构化内容供 AI Agent 使用。

支持格式:
  - .docx  (Word 文档: 文本/表格/图片/页眉/页脚/批注)
  - .pdf   (PDF 文档: 文本提取)
  - .txt   (纯文本)
  - .md    (Markdown)

安全性审计:
  ✅ 仅读取文件，不执行代码
  ✅ 文件大小限制防止内存溢出 (50MB)
  ✅ 图片 Base64 编码限制 (5张)
  ✅ 输入路径验证
"""

import os
import re
import io
import base64
import zipfile
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

# ---------- 最大限制常量 ----------
MAX_FILE_SIZE_MB = 50          # 最大文件大小 (MB)
MAX_IMAGE_COUNT = 5            # 最多提取图片数
MAX_TEXT_LENGTH = 50000        # 最大文本长度 (字符)


# ---------- 解析结果数据类 ----------
@dataclass
class ParsedDocument:
    """文档解析结果。

    Attributes:
        filename: 原始文件名
        format: 文件格式 (docx/pdf/txt/md)
        text: 提取的纯文本内容
        tables: 提取的表格列表 (每个表格为二维列表)
        images_base64: 图片 Base64 编码列表
        metadata: 元数据 (页数、作者等)
        headers: 页眉内容
        footers: 页脚内容
        comments: 批注内容
        errors: 解析过程中的警告/错误
    """
    filename: str = ""
    format: str = ""
    text: str = ""
    tables: List[List[List[str]]] = field(default_factory=list)
    images_base64: List[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    headers: List[str] = field(default_factory=list)
    footers: List[str] = field(default_factory=list)
    comments: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)

    def to_prompt_text(self) -> str:
        """将解析结果格式化为 LLM Prompt 友好的纯文本。

        Returns:
            格式化后的文本, 包含所有解析到的内容
        """
        parts = []
        parts.append(f"📄 文档: {self.filename} ({self.format.upper()})")
        parts.append("=" * 60)

        # 元数据
        if self.metadata:
            meta_items = [f"  {k}: {v}" for k, v in self.metadata.items() if v]
            if meta_items:
                parts.append("【文档信息】")
                parts.extend(meta_items)
                parts.append("")

        # 正文内容
        if self.text:
            text = self.text
            if len(text) > MAX_TEXT_LENGTH:
                text = text[:MAX_TEXT_LENGTH] + f"\n\n... (文本已截断, 共 {len(self.text)} 字符)"
            parts.append("【正文内容】")
            parts.append(text)
            parts.append("")

        # 表格
        if self.tables:
            parts.append(f"【表格内容】共 {len(self.tables)} 个表格")
            for i, table in enumerate(self.tables, 1):
                parts.append(f"\n--- 表格 {i} ---")
                for row in table:
                    parts.append(" | ".join(cell for cell in row))
            parts.append("")

        # 页眉
        if self.headers:
            unique_headers = list(set(h for h in self.headers if h.strip()))
            if unique_headers:
                parts.append("【页眉】")
                for h in unique_headers:
                    parts.append(f"  {h}")
                parts.append("")

        # 页脚
        if self.footers:
            unique_footers = list(set(f for f in self.footers if f.strip()))
            if unique_footers:
                parts.append("【页脚】")
                for f in unique_footers:
                    parts.append(f"  {f}")
                parts.append("")

        # 批注
        if self.comments:
            parts.append(f"【批注】共 {len(self.comments)} 条")
            for c in self.comments:
                parts.append(f"  💬 {c}")
            parts.append("")

        # 图片说明
        if self.images_base64:
            parts.append(f"【图片】共检测到 {len(self.images_base64)} 张图片 (已提取)")
            parts.append("")

        # 错误/警告
        if self.errors:
            parts.append("【解析警告】")
            for e in self.errors:
                parts.append(f"  ⚠️ {e}")

        return "\n".join(parts)


# ---------- 通用入口 ----------
def parse_document(file_path: str) -> ParsedDocument:
    """解析文档文件，自动识别格式。

    Args:
        file_path: 文件绝对路径

    Returns:
        ParsedDocument 解析结果对象
    """
    result = ParsedDocument()

    # 1. 验证文件存在
    if not os.path.exists(file_path):
        result.errors.append(f"文件不存在: {file_path}")
        return result

    # 2. 检查文件大小
    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE_MB:
        result.errors.append(
            f"文件过大 ({file_size_mb:.1f}MB > {MAX_FILE_SIZE_MB}MB 限制)")
        return result

    result.filename = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    # 3. 根据扩展名分发解析
    parsers = {
        ".docx": _parse_docx,
        ".pdf": _parse_pdf,
        ".txt": _parse_text,
        ".md": _parse_text,
        ".markdown": _parse_text,
        ".csv": _parse_text,
        ".json": _parse_text,
        ".xml": _parse_text,
        ".html": _parse_text,
        ".htm": _parse_text,
        ".log": _parse_text,
        ".py": _parse_text,
        ".js": _parse_text,
        ".yaml": _parse_text,
        ".yml": _parse_text,
        ".toml": _parse_text,
        ".ini": _parse_text,
        ".cfg": _parse_text,
    }

    parser_fn = parsers.get(ext)
    if parser_fn is None:
        result.errors.append(f"不支持的文件格式: {ext}")
        result.format = ext.lstrip(".")
        return result

    result.format = ext.lstrip(".")

    try:
        parser_fn(file_path, result)
    except Exception as e:
        result.errors.append(f"解析异常: {str(e)}")

    return result


# ==============================================================================
#  DOCX 解析 — 全内容类型
# ==============================================================================

def _parse_docx(file_path: str, result: ParsedDocument):
    """深度解析 Word 文档：文本/表格/图片/页眉/页脚/批注/元数据。"""
    try:
        import docx
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except ImportError:
        result.errors.append("python-docx 未安装。请运行: pip install python-docx")
        # 降级: 尝试纯文本提取
        _parse_docx_fallback(file_path, result)
        return

    doc = docx.Document(file_path)

    # --- 1. 元数据 ---
    props = doc.core_properties
    result.metadata = {
        "标题": props.title or "",
        "作者": props.author or "",
        "创建时间": str(props.created) if props.created else "",
        "修改时间": str(props.modified) if props.modified else "",
        "修订版本": str(props.revision) if props.revision else "",
        "主题": props.subject or "",
        "关键词": props.keywords or "",
        "描述": props.comments or "",
    }

    # --- 2. 正文内容 (段落 + 内联图片标记) ---
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 检测标题级别
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                try:
                    level_int = int(level)
                except ValueError:
                    level_int = 1
                prefix = "#" * level_int + " "
                paragraphs.append(f"{prefix}{text}")
            else:
                paragraphs.append(text)

    result.text = "\n\n".join(paragraphs)

    # --- 3. 表格 ---
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # 合并单元格可能重复，取纯文本
                cell_text = cell.text.strip().replace("\n", " ")
                row_data.append(cell_text)
            table_data.append(row_data)
        if table_data:
            result.tables.append(table_data)

    # --- 4. 页眉 ---
    for section in doc.sections:
        header = section.header
        if header and not header.is_linked_to_previous:
            header_text = "\n".join(
                p.text.strip() for p in header.paragraphs if p.text.strip()
            )
            if header_text:
                result.headers.append(header_text)

    # --- 5. 页脚 ---
    for section in doc.sections:
        footer = section.footer
        if footer and not footer.is_linked_to_previous:
            footer_text = "\n".join(
                p.text.strip() for p in footer.paragraphs if p.text.strip()
            )
            if footer_text:
                result.footers.append(footer_text)

    # --- 6. 批注 (Comments) ---
    try:
        _extract_comments(file_path, result)
    except Exception as e:
        result.errors.append(f"批注提取失败: {e}")

    # --- 7. 图片 ---
    try:
        _extract_images(file_path, result)
    except Exception as e:
        result.errors.append(f"图片提取失败: {e}")


def _parse_docx_fallback(file_path: str, result: ParsedDocument):
    """降级解析: 不依赖 python-docx, 直接解压 XML 提取文本。"""
    try:
        with zipfile.ZipFile(file_path) as z:
            if "word/document.xml" in z.namelist():
                xml_content = z.read("word/document.xml").decode("utf-8")
                # 简单正则提取 <w:t> 标签中的文本
                texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml_content)
                result.text = " ".join(texts)
                result.errors.append("使用降级模式解析 (仅提取纯文本)")
    except Exception as e:
        result.errors.append(f"降级解析也失败: {e}")


def _extract_comments(file_path: str, result: ParsedDocument):
    """从 docx ZIP 中提取批注内容。"""
    with zipfile.ZipFile(file_path) as z:
        if "word/comments.xml" in z.namelist():
            xml_content = z.read("word/comments.xml").decode("utf-8")
            # 提取批注文本
            comment_texts = re.findall(
                r"<w:comment\b[^>]*>.*?<w:t[^>]*>(.*?)</w:t>.*?</w:comment>",
                xml_content,
                re.DOTALL,
            )
            for ct in comment_texts:
                ct_clean = ct.strip()
                if ct_clean:
                    result.comments.append(ct_clean)


def _extract_images(file_path: str, result: ParsedDocument):
    """从 docx ZIP 中提取嵌入图片并转为 Base64。"""
    image_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')
    count = 0

    with zipfile.ZipFile(file_path) as z:
        for name in z.namelist():
            if name.startswith("word/media/") and name.lower().endswith(image_extensions):
                if count >= MAX_IMAGE_COUNT:
                    result.errors.append(
                        f"图片数量超过限制 ({MAX_IMAGE_COUNT} 张), 后续图片已跳过")
                    break
                try:
                    img_data = z.read(name)
                    # 验证是否为合法图片
                    from PIL import Image
                    img = Image.open(io.BytesIO(img_data))
                    img.verify()  # 验证图片完整性

                    # 压缩大图
                    img = Image.open(io.BytesIO(img_data))
                    if img.width > 800 or img.height > 800:
                        img.thumbnail((800, 800), Image.Resampling.LANCZOS)

                    # 转 Base64
                    buffer = io.BytesIO()
                    img_format = "PNG" if name.lower().endswith(".png") else "JPEG"
                    img.save(buffer, format=img_format)
                    b64 = base64.b64encode(buffer.getvalue()).decode("utf-8")
                    result.images_base64.append(
                        f"data:image/{img_format.lower()};base64,{b64}"
                    )
                    count += 1
                except ImportError:
                    result.errors.append("Pillow 未安装, 图片提取已跳过")
                    break
                except Exception as e:
                    result.errors.append(f"图片 {os.path.basename(name)} 解析失败: {e}")


# ==============================================================================
#  PDF 解析
# ==============================================================================

def _parse_pdf(file_path: str, result: ParsedDocument):
    """解析 PDF 文档, 提取文本内容。"""
    # 尝试多个 PDF 库
    parsed = False

    # 方案 1: PyMuPDF (fitz)
    if not parsed:
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            result.metadata["页数"] = str(len(doc))

            pages_text = []
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text("text")
                if page_text.strip():
                    pages_text.append(f"--- 第 {page_num} 页 ---\n{page_text.strip()}")

            result.text = "\n\n".join(pages_text)
            doc.close()
            parsed = True
        except ImportError:
            pass

    # 方案 2: pdfplumber
    if not parsed:
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                result.metadata["页数"] = str(len(pdf.pages))
                pages_text = []
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        pages_text.append(f"--- 第 {i} 页 ---\n{text.strip()}")
                    # 提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        result.tables.append(
                            [[cell or "" for cell in row] for row in table]
                        )
                result.text = "\n\n".join(pages_text)
            parsed = True
        except ImportError:
            pass

    # 方案 3: PyPDF2
    if not parsed:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            result.metadata["页数"] = str(len(reader.pages))

            pages_text = []
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append(f"--- 第 {i} 页 ---\n{text.strip()}")
            result.text = "\n\n".join(pages_text)
            parsed = True
        except ImportError:
            pass

    if not parsed:
        result.errors.append(
            "PDF 解析需要安装以下库之一: "
            "pip install PyMuPDF 或 pip install pdfplumber 或 pip install PyPDF2"
        )


# ==============================================================================
#  纯文本 / Markdown 解析
# ==============================================================================

def _parse_text(file_path: str, result: ParsedDocument):
    """读取纯文本文件 (txt/md/csv/json 等)。"""
    # 尝试多种编码
    encodings = ["utf-8", "gbk", "gb2312", "utf-16", "latin-1"]

    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                result.text = f.read()
            result.metadata["编码"] = encoding
            result.metadata["行数"] = str(result.text.count("\n") + 1)
            return
        except (UnicodeDecodeError, UnicodeError):
            continue

    result.errors.append(f"无法以任何已知编码读取文件 (尝试: {', '.join(encodings)})")
