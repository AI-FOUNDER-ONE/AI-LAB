"""
debug_word_io.py - Test Word Document Cloning (Full Loop)
=========================================================
"""
import os
import sys
import logging
from agents.coder_agent import CoderAgent
from PyQt6.QtCore import QCoreApplication

# Setup Qt App
app = QCoreApplication(sys.argv)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_mock_docx(filename, content_lines):
    try:
        import docx
        doc = docx.Document()
        doc.add_heading("Mock Document", 0)
        for line in content_lines:
            doc.add_paragraph(line)
        doc.save(filename)
        logger.info(f"Created mock doc: {filename}")
        return True
    except ImportError:
        logger.error("python-docx not installed")
        return False

def main():
    # 1. Create Mock Files
    proto_file = "mock_prototype.docx"
    content_file = "mock_content.docx"
    
    if not create_mock_docx(proto_file, [
        "1. Project Background", 
        "This section describes the project background.",
        "2. Technical Architecture",
        "This section describes the technical details."
    ]):
        return

    create_mock_docx(content_file, [
        "We need to build a new AI system.",
        "The background is that our old system is too slow.",
        "Technically, we want to use Python and TensorFlow."
    ])

    # 2. Initialize Coder
    logger.info("Initializing CoderAgent...")
    coder = CoderAgent()
    
    # 3. Run Cloning with File Paths
    logger.info("Running Document Cloning with File Paths...")
    try:
        # Pass file paths directly
        result = coder.run_document_cloning(
            os.path.abspath(proto_file), 
            os.path.abspath(content_file)
        )
        print("\n--- Result ---")
        print(result)
        print("----------------")
    except Exception as e:
        logger.error(f"Failed: {e}")

if __name__ == "__main__":
    main()
